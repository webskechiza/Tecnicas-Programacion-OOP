"""
Script para generar VetCare_Avance2.docx
Diseño: verde esmeralda + ámbar dorado (paleta profesional veterinaria)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paleta de colores ──────────────────────────────────────────────────────────
VERDE_OSCURO   = RGBColor(0x14, 0x5A, 0x32)   # #145A32 verde oscuro
VERDE_MEDIO    = RGBColor(0x1E, 0x8B, 0x4F)   # #1E8B4F verde medio
VERDE_CLARO    = RGBColor(0xD5, 0xF5, 0xE3)   # #D5F5E3 verde muy claro (fondo tabla)
AMBAR          = RGBColor(0xB7, 0x7A, 0x0C)   # #B77A0C ámbar dorado
AMBAR_CLARO    = RGBColor(0xFD, 0xF2, 0xD0)   # #FDF2D0 ámbar muy claro
BLANCO         = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO     = RGBColor(0x2C, 0x3E, 0x50)   # #2C3E50 gris oscuro texto
GRIS_CLARO     = RGBColor(0xF4, 0xF6, 0xF7)   # #F4F6F7 fondo alternativo

def set_cell_color(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '6')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_with_border(doc, text, color_hex, font_color=None, size=11, bold=False, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if font_color:
        run.font.color.rgb = font_color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Borde izquierdo decorativo
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '18')
    left_border.set(qn('w:space'), '12')
    left_border.set(qn('w:color'), color_hex)
    pBdr.append(left_border)
    pPr.append(pBdr)
    # Fondo del párrafo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F8FDF9')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    return p

def add_section_title(doc, numero, titulo):
    p = doc.add_paragraph()
    p.clear()
    # Fondo verde oscuro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '145A32')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    # Spacing
    pPr2 = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '120')
    pPr2.append(spacing)
    # Número
    run_num = p.add_run(f"  {numero}.   ")
    run_num.font.size = Pt(13)
    run_num.font.bold = True
    run_num.font.color.rgb = AMBAR
    # Título
    run_title = p.add_run(titulo.upper())
    run_title.font.size = Pt(13)
    run_title.font.bold = True
    run_title.font.color.rgb = BLANCO
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_subsection_title(doc, titulo):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E8F8F0')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(f"  {titulo}")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = VERDE_OSCURO
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTO
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, color_bullet='145A32'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    # Bullet verde
    run_bullet = p.add_run("▸  ")
    run_bullet.font.color.rgb = VERDE_MEDIO
    run_bullet.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TEXTO
    return p

def add_bullet_red(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run_bullet = p.add_run("✗  ")
    run_bullet.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run_bullet.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TEXTO
    return p

def add_alternativa(doc, restriccion, alternativa):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Fila 1: restricción
    cell1 = tbl.rows[0].cells[0]
    cell1.text = ""
    p1 = cell1.paragraphs[0]
    run_label = p1.add_run("Restricción:  ")
    run_label.font.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run_rest = p1.add_run(restriccion)
    run_rest.font.size = Pt(10)
    run_rest.font.color.rgb = GRIS_TEXTO
    set_cell_color(cell1, 'FEF9E7')
    cell1._tc.get_or_add_tcPr()
    # Fila 2: alternativa
    cell2 = tbl.rows[1].cells[0]
    cell2.text = ""
    p2 = cell2.paragraphs[0]
    run_label2 = p2.add_run("Alternativa de solución:  ")
    run_label2.font.bold = True
    run_label2.font.size = Pt(10)
    run_label2.font.color.rgb = VERDE_OSCURO
    run_alt = p2.add_run(alternativa)
    run_alt.font.size = Pt(10)
    run_alt.font.color.rgb = GRIS_TEXTO
    set_cell_color(cell2, 'EAFAF1')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_rf_table(doc, modulo, rfs):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E8B4F')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(f"  {modulo}")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLANCO

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    for i, txt in enumerate(['ID', 'Nombre', 'Descripción', 'Prioridad']):
        cell = hdr.cells[i]
        cell.text = txt
        run_h = cell.paragraphs[0].runs[0]
        run_h.font.bold = True
        run_h.font.size = Pt(9.5)
        run_h.font.color.rgb = BLANCO
        set_cell_color(cell, '145A32')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (rid, nombre, desc, prio) in enumerate(rfs):
        row = tbl.add_row()
        data = [rid, nombre, desc, prio]
        bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
        for i, txt in enumerate(data):
            row.cells[i].text = txt
            run_d = row.cells[i].paragraphs[0].runs[0]
            run_d.font.size = Pt(9)
            run_d.font.color.rgb = GRIS_TEXTO
            set_cell_color(row.cells[i], bg)
        # Color prioridad
        prio_cell = row.cells[3]
        prio_run = prio_cell.paragraphs[0].runs[0]
        if prio == 'Alta':
            prio_run.font.color.rgb = RGBColor(0x1A, 0x5E, 0x20)
            prio_run.font.bold = True
        elif prio == 'Media':
            prio_run.font.color.rgb = RGBColor(0x7D, 0x60, 0x08)
        elif prio == 'Baja':
            prio_run.font.color.rgb = RGBColor(0x78, 0x2F, 0x0E)
        prio_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ancho de columnas
    for i, w in enumerate([Cm(1.5), Cm(3.5), Cm(9.0), Cm(2.0)]):
        for cell in tbl.columns[i].cells:
            cell.width = w

def add_hu_card(doc, hu_id, titulo, rol, quiero, para, rfs, escenario_ok, escenario_fail):
    # Título de la tarjeta
    p_title = doc.add_paragraph()
    pPr = p_title._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '145A32')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run_id = p_title.add_run(f"  {hu_id}  ·  ")
    run_id.font.color.rgb = AMBAR
    run_id.font.bold = True
    run_id.font.size = Pt(10.5)
    run_t = p_title.add_run(titulo)
    run_t.font.color.rgb = BLANCO
    run_t.font.bold = True
    run_t.font.size = Pt(10.5)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    labels = ['Como', 'Quiero', 'Para', 'RF relacionados', 'Escenario exitoso']
    values = [rol, quiero, para, rfs, escenario_ok]
    for i, (label, value) in enumerate(zip(labels, values)):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = VERDE_OSCURO
        set_cell_color(tbl.rows[i].cells[0], 'E8F8F0')
        tbl.rows[i].cells[1].text = value
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = GRIS_TEXTO
        set_cell_color(tbl.rows[i].cells[1], 'FFFFFF')

    if escenario_fail:
        row_fail = tbl.add_row()
        row_fail.cells[0].text = 'Escenario fallido'
        row_fail.cells[0].paragraphs[0].runs[0].font.bold = True
        row_fail.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_fail.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAB, 0x27, 0x12)
        set_cell_color(row_fail.cells[0], 'FDEDEC')
        row_fail.cells[1].text = escenario_fail
        row_fail.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_color(row_fail.cells[1], 'FEF5F5')

    # Ancho columnas
    for cell in tbl.columns[0].cells:
        cell.width = Cm(3.5)
    for cell in tbl.columns[1].cells:
        cell.width = Cm(12.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Fuente por defecto
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── PORTADA ───────────────────────────────────────────────────────────────────
# Header institucional verde
p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_inst._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), '145A32')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:before'), '160')
spacing.set(qn('w:after'), '160')
pPr.append(spacing)
run = p_inst.add_run("UNIVERSIDAD PRIVADA DEL NORTE")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = BLANCO

p_fac = doc.add_paragraph()
p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2 = p_fac._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:fill'), '1E8B4F')
shd2.set(qn('w:val'), 'clear')
pPr2.append(shd2)
r1 = p_fac.add_run("Facultad de Ingeniería  ·  Técnicas de Programación Orientada a Objetos")
r1.font.size = Pt(11)
r1.font.color.rgb = BLANCO

doc.add_paragraph()

# Nombre del sistema con efecto visual
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr3 = p_name._p.get_or_add_pPr()
shd3 = OxmlElement('w:shd')
shd3.set(qn('w:fill'), 'EAFAF1')
shd3.set(qn('w:val'), 'clear')
pPr3.append(shd3)
spacing3 = OxmlElement('w:spacing')
spacing3.set(qn('w:before'), '240')
spacing3.set(qn('w:after'), '80')
pPr3.append(spacing3)
r_name = p_name.add_run("🐾  VETCARE")
r_name.font.size = Pt(32)
r_name.font.bold = True
r_name.font.color.rgb = VERDE_OSCURO

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr4 = p_sub._p.get_or_add_pPr()
shd4 = OxmlElement('w:shd')
shd4.set(qn('w:fill'), 'EAFAF1')
shd4.set(qn('w:val'), 'clear')
pPr4.append(shd4)
r_sub = p_sub.add_run("Sistema de Gestión de Clínica Veterinaria")
r_sub.font.size = Pt(16)
r_sub.font.color.rgb = AMBAR
r_sub.font.bold = True

p_avance = doc.add_paragraph()
p_avance.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr5 = p_avance._p.get_or_add_pPr()
shd5 = OxmlElement('w:shd')
shd5.set(qn('w:fill'), 'EAFAF1')
shd5.set(qn('w:val'), 'clear')
pPr5.append(shd5)
spacing5 = OxmlElement('w:spacing')
spacing5.set(qn('w:before'), '80')
spacing5.set(qn('w:after'), '240')
pPr5.append(spacing5)
r_av = p_avance.add_run("AVANCE 2 — Práctica de Campo")
r_av.font.size = Pt(13)
r_av.font.color.rgb = VERDE_MEDIO
r_av.font.bold = True

doc.add_paragraph()

# Tabla de datos del alumno
tbl_info = doc.add_table(rows=3, cols=2)
tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER
datos = [
    ('Integrante:', 'Kevin Oswaldo Chirinos Zapata – N00521954'),
    ('Docente:', 'Martín Eduardo Torres Rodríguez'),
    ('Año:', '2026'),
]
for i, (label, value) in enumerate(datos):
    tbl_info.rows[i].cells[0].text = label
    tbl_info.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    tbl_info.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = VERDE_OSCURO
    tbl_info.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10.5)
    set_cell_color(tbl_info.rows[i].cells[0], 'D5F5E3')
    tbl_info.rows[i].cells[1].text = value
    tbl_info.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
    tbl_info.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = GRIS_TEXTO
    set_cell_color(tbl_info.rows[i].cells[1], 'F8FDF9')
for cell in tbl_info.columns[0].cells:
    cell.width = Cm(4)
for cell in tbl_info.columns[1].cells:
    cell.width = Cm(10)

doc.add_paragraph()

# Línea dorada decorativa
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_l = p_line._p.get_or_add_pPr()
shd_l = OxmlElement('w:shd')
shd_l.set(qn('w:fill'), 'B77A0C')
shd_l.set(qn('w:val'), 'clear')
pPr_l.append(shd_l)
spacing_l = OxmlElement('w:spacing')
spacing_l.set(qn('w:before'), '0')
spacing_l.set(qn('w:after'), '0')
pPr_l.append(spacing_l)
r_line = p_line.add_run(" " * 80)
r_line.font.size = Pt(4)

doc.add_page_break()

# ── ÍNDICE ────────────────────────────────────────────────────────────────────
p_idx_title = doc.add_paragraph()
p_idx_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_idx = p_idx_title._p.get_or_add_pPr()
shd_idx = OxmlElement('w:shd')
shd_idx.set(qn('w:fill'), '145A32')
shd_idx.set(qn('w:val'), 'clear')
pPr_idx.append(shd_idx)
r_idx = p_idx_title.add_run("  ÍNDICE DE CONTENIDOS  ")
r_idx.font.size = Pt(14)
r_idx.font.bold = True
r_idx.font.color.rgb = BLANCO

indice_items = [
    ("Introducción", "3"),
    ("1.  Definición del Problema", "4"),
    ("2.  Antecedentes", "5"),
    ("3.  Restricciones Realistas y Alternativas de Solución", "6"),
    ("4.  Objetivos del Proyecto", "8"),
    ("5.  Alcance de la Solución", "9"),
    ("6.  Requerimientos Funcionales", "10"),
    ("7.  Historias de Usuario", "14"),
    ("8.  Diagrama de Clases", "19"),
    ("9.  Criterios de Aceptación", "20"),
]
for item, page in indice_items:
    p_i = doc.add_paragraph()
    p_i.paragraph_format.left_indent = Cm(1.0)
    p_i.paragraph_format.space_after = Pt(3)
    tab = p_i.add_run(item)
    tab.font.size = Pt(10.5)
    tab.font.color.rgb = GRIS_TEXTO
    dots = p_i.add_run("  " + "·" * (60 - len(item)) + "  " + page)
    dots.font.size = Pt(10.5)
    dots.font.color.rgb = VERDE_MEDIO

doc.add_page_break()

# ── INTRODUCCIÓN ──────────────────────────────────────────────────────────────
p_intro_title = doc.add_paragraph()
pPr_int = p_intro_title._p.get_or_add_pPr()
shd_int = OxmlElement('w:shd')
shd_int.set(qn('w:fill'), '145A32')
shd_int.set(qn('w:val'), 'clear')
pPr_int.append(shd_int)
r_int = p_intro_title.add_run("  INTRODUCCIÓN")
r_int.font.size = Pt(13)
r_int.font.bold = True
r_int.font.color.rgb = BLANCO

add_body(doc, "VetCare es un sistema de gestión diseñado para digitalizar y optimizar los procesos internos "
    "de las clínicas veterinarias de pequeño y mediano tamaño. Este proyecto surge de la necesidad urgente "
    "de modernizar la gestión de pacientes animales, la programación de citas médicas y el mantenimiento "
    "de historiales clínicos que, en la mayoría de los establecimientos veterinarios del Perú, aún se "
    "realizan de forma manual y desorganizada.", indent=False)

add_body(doc, "El sistema ha sido desarrollado en Java, aplicando los principios de la Programación Orientada "
    "a Objetos (POO). Este enfoque permite representar las entidades del negocio veterinario (mascotas, "
    "dueños, veterinarios y citas) como clases con atributos y métodos propios, garantizando un código "
    "modular, reutilizable y fácil de mantener.", indent=False)

add_body(doc, "A lo largo de este documento se presentan los fundamentos del proyecto: la definición del "
    "problema central, los antecedentes que respaldan la propuesta, las restricciones técnicas y operativas "
    "con sus respectivas alternativas de solución, los objetivos del proyecto, el alcance funcional, los 40 "
    "requerimientos del sistema, las 10 historias de usuario, el diagrama de clases y los criterios de "
    "aceptación.", indent=False)

doc.add_page_break()

# ── SECCIÓN 1: DEFINICIÓN DEL PROBLEMA ───────────────────────────────────────
add_section_title(doc, 1, "Definición del Problema")
doc.add_paragraph()

add_subsection_title(doc, "1.1  Situación Problemática")
add_body(doc, "Las clínicas veterinarias de pequeño y mediano tamaño en el Perú operan, en su mayoría, "
    "con procesos manuales que limitan su eficiencia operativa. El registro de mascotas, la programación "
    "de citas y el mantenimiento del historial médico se realizan en cuadernos físicos o planillas de "
    "Excel, lo que genera duplicidad de datos, pérdida de información y demoras en la atención.")

add_body(doc, "Esta situación se agrava en períodos de alta demanda (campañas de vacunación o desparasitación "
    "masiva) donde la ausencia de un sistema organizado de citas provoca conflictos de horario y una atención "
    "deficiente. Asimismo, la falta de un historial médico digital centralizado impide al veterinario "
    "acceder rápidamente al historial clínico del paciente durante la consulta.")

add_subsection_title(doc, "1.2  Problema Identificado")
p_q = doc.add_paragraph()
p_q.paragraph_format.left_indent = Cm(0.5)
pPr_q = p_q._p.get_or_add_pPr()
shd_q = OxmlElement('w:shd')
shd_q.set(qn('w:fill'), 'EAFAF1')
shd_q.set(qn('w:val'), 'clear')
pPr_q.append(shd_q)
r_q = p_q.add_run("¿De qué manera el desarrollo de un sistema de gestión de clínica veterinaria en Java "
    "orientado a objetos puede digitalizar los procesos de registro de mascotas, programación de citas e "
    "historial médico, mejorando la eficiencia operativa y la calidad de atención de la clínica?")
r_q.font.size = Pt(10.5)
r_q.font.italic = True
r_q.font.color.rgb = VERDE_OSCURO

add_subsection_title(doc, "1.3  Diagrama de Ishikawa")
add_body(doc, 'Problema central: "Ineficiencia en la gestión de atención de pacientes en clínicas veterinarias de pequeño tamaño"')

ishikawa_categorias = {
    "Personal:":         ["Falta de capacitación tecnológica del personal",
                          "Resistencia al cambio y adopción de herramientas digitales",
                          "Escaso personal para atención y administración simultánea"],
    "Método:":           ["Ausencia de protocolos de registro estandarizados",
                          "Programación de citas sin sistema de confirmación",
                          "Sin proceso de seguimiento post-consulta al dueño"],
    "Tecnología:":       ["Ausencia de software especializado para veterinarias",
                          "Dependencia de cuadernos físicos y hojas de cálculo",
                          "Sin acceso a historial médico digital en tiempo real"],
    "Dueños:":           ["Sin recordatorios automáticos de citas o vacunas",
                          "Dificultad para consultar el historial de su mascota",
                          "Sin canal de comunicación directo con la clínica"],
    "Historial médico:": ["Registros incompletos o ilegibles en papel",
                          "Información dispersa y sin centralización",
                          "Riesgo de pérdida de datos ante extravíos o siniestros"],
    "Finanzas:":         ["Cobros manuales sin registro sistemático",
                          "Sin reportes de ingresos por período",
                          "Dificultad para realizar cierres de caja"],
}
for categoria, items in ishikawa_categorias.items():
    add_subsection_title(doc, f"  {categoria}")
    for item in items:
        add_bullet(doc, item)

add_body(doc, "El diagrama completo de causa-efecto se presenta a continuación:")

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
img_path = os.path.join(SCRIPT_DIR, "Semana3", "VetCare_Documentos", "ishikawa_vetcare.png")
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Cm(14))
    p_caption = doc.add_paragraph("Figura 1. Diagrama de Ishikawa — VetCare")
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.runs[0].font.size = Pt(9)
    p_caption.runs[0].font.italic = True
    p_caption.runs[0].font.color.rgb = RGBColor(0x64, 0x64, 0x64)

doc.add_page_break()

# ── SECCIÓN 2: ANTECEDENTES ───────────────────────────────────────────────────
add_section_title(doc, 2, "Antecedentes")
doc.add_paragraph()
add_body(doc, "A continuación se presentan tres antecedentes que respaldan el desarrollo de VetCare "
    "desde perspectivas internacional, nacional y tecnológica.", indent=False)

antecedentes = [
    ("Antecedente Internacional",
     "Barrientos, J., & Molina, P. (2020). Sistema de información para la gestión de clínicas veterinarias. "
     "Repositorio Universidad de Guayaquil, Ecuador.",
     "Este trabajo implementó un sistema de escritorio en Java para automatizar el registro de pacientes "
     "animales y la gestión de citas en clínicas veterinarias de Guayaquil, Ecuador. Los resultados evidenciaron "
     "una reducción del 70% en el tiempo de registro de pacientes y una mejora significativa en la organización "
     "del historial médico. El proyecto guarda similitud directa con VetCare en cuanto al uso de Java y la "
     "gestión de citas como módulo central."),
    ("Antecedente Nacional",
     "Quispe, A., & Vargas, L. (2023). Desarrollo de software de escritorio para la gestión de consultas "
     "en clínicas veterinarias de Lima Metropolitana. Repositorio Institucional USIL, Lima, Perú.",
     "Este estudio desarrolló una aplicación de escritorio para la clínica veterinaria 'Patitas Felices' "
     "en el distrito de San Borja, Lima. El sistema implementado redujo los errores en el registro de citas "
     "en un 85% y permitió centralizar el historial médico de más de 500 mascotas. Los autores destacan que "
     "el uso de POO facilitó la escalabilidad del sistema, alineándose con la propuesta arquitectónica de "
     "VetCare."),
    ("Antecedente Tecnológico",
     "Deitel, P. J., & Deitel, H. M. (2012). Java: Cómo programar (9.a ed.). Pearson Educación.",
     "Esta referencia canónica del lenguaje Java fundamenta el paradigma de la Programación Orientada a "
     "Objetos como base para el desarrollo de sistemas de gestión. Los autores demuestran que la representación "
     "de entidades del mundo real (mascotas, médicos, citas) como clases con atributos y métodos propios "
     "constituye la estrategia óptima para sistemas de escritorio escalables y mantenibles."),
]
for titulo, ref, desc in antecedentes:
    add_subsection_title(doc, titulo)
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.left_indent = Cm(0.5)
    run_ref = p_ref.add_run(ref)
    run_ref.font.size = Pt(10)
    run_ref.font.italic = True
    run_ref.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    add_body(doc, desc)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── SECCIÓN 3: RESTRICCIONES REALISTAS Y ALTERNATIVAS ────────────────────────
add_section_title(doc, 3, "Restricciones Realistas y Alternativas de Solución")
doc.add_paragraph()
add_body(doc, "A continuación se identifican las restricciones que afectan el desarrollo del proyecto VetCare, "
    "organizadas por categoría, junto con las alternativas de solución propuestas para cada caso.", indent=False)

add_subsection_title(doc, "3.1  Restricciones Técnicas")

rest_tecnicas = [
    ("El sistema se desarrolla en Java SE (JDK 17 o superior), sin uso de frameworks externos.",
     "Se aprovechan al máximo las librerías estándar de Java (java.util, java.io, java.time). "
     "En fases futuras se podría incorporar JDBC para acceso a base de datos sin cambiar el lenguaje."),
    ("La persistencia de datos se realiza mediante archivos de texto (.txt), sin base de datos relacional.",
     "Se implementa un gestor de archivos robusto con manejo de errores (try-catch) y separadores "
     "de campo definidos. En versiones posteriores se podría migrar a SQLite embebido sin cambiar la arquitectura."),
    ("La interfaz de usuario es de consola, sin interfaz gráfica (GUI/Swing/JavaFX).",
     "Se diseñan menús numerados intuitivos con validaciones en cada entrada y mensajes de error "
     "descriptivos, logrando una experiencia de usuario clara a pesar de la limitación gráfica."),
    ("El entorno de desarrollo es Eclipse IDE.",
     "El código es compatible con IntelliJ IDEA Community Edition y VS Code con extensión Java, "
     "lo que permite que cualquier integrante del equipo trabaje con su herramienta preferida."),
    ("El control de versiones se gestiona con Git y el repositorio se aloja en GitHub.",
     "En caso de inconvenientes de acceso, el repositorio puede duplicarse en GitLab como espejo "
     "de respaldo sin modificar el flujo de trabajo."),
]
for r, a in rest_tecnicas:
    add_alternativa(doc, r, a)

add_subsection_title(doc, "3.2  Restricciones Económicas")

rest_economicas = [
    ("Se utilizan exclusivamente herramientas gratuitas: Java SE, Eclipse IDE, Git y GitHub.",
     "Las versiones Community Edition de IntelliJ IDEA ofrecen funciones avanzadas de refactorización "
     "sin costo. Oracle Cloud Free Tier también puede usarse para pruebas de despliegue sin presupuesto."),
    ("No se cuenta con presupuesto para licencias de software, servidores ni servicios en la nube.",
     "GitHub provee repositorios públicos y GitHub Actions (CI) gratuitos, cubriendo las necesidades "
     "de control de versiones y automatización sin costo alguno."),
    ("El proyecto se desarrolla con los equipos personales de los integrantes del equipo.",
     "Los laboratorios de cómputo de la universidad están disponibles para sesiones colaborativas "
     "presenciales cuando el equipo personal no esté disponible."),
]
for r, a in rest_economicas:
    add_alternativa(doc, r, a)

add_subsection_title(doc, "3.3  Restricciones Operativas")

rest_operativas = [
    ("El tiempo de desarrollo está limitado al ciclo académico 2026, con entregas parciales semanales.",
     "Se aplica una estrategia de priorización por módulos (MVP): primero mascotas, citas e historial; "
     "el módulo de reportes y autenticación se integra en la fase final del ciclo."),
    ("Los integrantes tienen carga académica paralela, lo que limita la dedicación exclusiva al proyecto.",
     "Se establecen reuniones semanales cortas de sincronización y se usa comunicación asíncrona "
     "(WhatsApp/Discord) para resolver bloqueos sin necesidad de reuniones adicionales."),
    ("El sistema no incluye integración con sistemas externos (SUNAT, pasarelas de pago, APIs de terceros).",
     "El diseño de clases contempla interfaces preparadas para futura integración REST, de modo que "
     "agregar conectores externos en versiones posteriores no requiera refactorización mayor."),
    ("El módulo de recordatorios automáticos por correo o SMS queda fuera del alcance del ciclo.",
     "Se implementa una función de 'agenda del día' que el recepcionista consulta al iniciar la jornada, "
     "cubriendo operativamente la necesidad de recordatorios de forma manual."),
]
for r, a in rest_operativas:
    add_alternativa(doc, r, a)

doc.add_page_break()

# ── SECCIÓN 4: OBJETIVOS ──────────────────────────────────────────────────────
add_section_title(doc, 4, "Objetivos del Proyecto")
doc.add_paragraph()

add_subsection_title(doc, "4.1  Objetivo General")
p_og = doc.add_paragraph()
p_og.paragraph_format.left_indent = Cm(0.5)
pPr_og = p_og._p.get_or_add_pPr()
shd_og = OxmlElement('w:shd')
shd_og.set(qn('w:fill'), 'EAFAF1')
shd_og.set(qn('w:val'), 'clear')
pPr_og.append(shd_og)
r_og = p_og.add_run('Desarrollar un sistema de gestión de clínica veterinaria denominado "VetCare" en Java '
    "orientado a objetos, que digitalice los procesos de registro de mascotas, gestión de citas médicas e "
    "historial clínico, mejorando la eficiencia operativa y la calidad de atención de la clínica.")
r_og.font.size = Pt(11)
r_og.font.color.rgb = GRIS_TEXTO

add_subsection_title(doc, "4.2  Objetivos Específicos")

oe_table = doc.add_table(rows=1, cols=2)
oe_table.style = 'Table Grid'
for hdr_txt, cell_obj in zip(['ID', 'Objetivo Específico'], oe_table.rows[0].cells):
    cell_obj.text = hdr_txt
    cell_obj.paragraphs[0].runs[0].font.bold = True
    cell_obj.paragraphs[0].runs[0].font.size = Pt(10)
    cell_obj.paragraphs[0].runs[0].font.color.rgb = BLANCO
    set_cell_color(cell_obj, '145A32')
    cell_obj.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

oes = [
    ("OE-01", "Implementar un módulo de gestión de mascotas que permita registrar, consultar y mantener "
              "el historial médico de cada paciente animal."),
    ("OE-02", "Desarrollar un módulo de gestión de citas que organice la programación de atenciones "
              "veterinarias con control de conflictos de horario."),
    ("OE-03", "Crear un módulo de gestión de clientes (dueños) que centralice la información de contacto "
              "y su vinculación con las mascotas registradas."),
    ("OE-04", "Diseñar un diagrama de clases UML que aplique herencia (Animal → Perro/Gato/Ave), "
              "composición y asociación entre las entidades del sistema."),
    ("OE-05", "Aplicar principios de POO (herencia, encapsulamiento, sobrecarga de métodos) y colecciones "
              "ArrayList para estructurar el código de forma modular y escalable."),
]
for idx, (oe_id, oe_desc) in enumerate(oes):
    row = oe_table.add_row()
    row.cells[0].text = oe_id
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = VERDE_OSCURO
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    set_cell_color(row.cells[0], bg)
    row.cells[1].text = oe_desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = GRIS_TEXTO
    set_cell_color(row.cells[1], bg)

for cell in oe_table.columns[0].cells:
    cell.width = Cm(2.0)
for cell in oe_table.columns[1].cells:
    cell.width = Cm(14.0)

doc.add_page_break()

# ── SECCIÓN 5: ALCANCE ────────────────────────────────────────────────────────
add_section_title(doc, 5, "Alcance de la Solución")
doc.add_paragraph()
add_body(doc, "El alcance define con precisión qué funcionalidades incluye VetCare en el ciclo 2026 "
    "y cuáles quedan fuera del desarrollo actual.", indent=False)

add_subsection_title(doc, "Funcionalidades incluidas")
incluidas = [
    "Módulo de autenticación con roles (administrador / recepcionista).",
    "Gestión completa de mascotas: registro, modificación, eliminación y búsqueda.",
    "Clasificación de mascotas por especie con herencia UML (Perro, Gato, Ave).",
    "Gestión de dueños: registro, búsqueda por nombre y DNI, historial de mascotas.",
    "Gestión de médicos veterinarios: registro, listado y agenda de citas.",
    "Gestión de citas: registro, modificación, cancelación y confirmación de asistencia.",
    "Registro de consultas médicas: diagnóstico, tratamiento y fecha.",
    "Historial médico digital por mascota.",
    "Reportes básicos de citas y consultas por período.",
    "Persistencia de datos en archivos de texto (.txt).",
]
for item in incluidas:
    add_bullet(doc, item)

add_subsection_title(doc, "Funcionalidades NO incluidas")
no_incluidas = [
    "Interfaz gráfica de usuario (GUI / Swing / JavaFX).",
    "Conexión a base de datos relacional (MySQL, PostgreSQL).",
    "Módulo de facturación, pagos o integración con POS.",
    "Envío de recordatorios automáticos por correo electrónico o SMS.",
    "Aplicación web o móvil.",
    "Módulo de farmacia o control de medicamentos.",
    "Integración con sistemas externos de identificación de mascotas.",
]
for item in no_incluidas:
    add_bullet_red(doc, item)

doc.add_page_break()

# ── SECCIÓN 6: REQUERIMIENTOS FUNCIONALES ─────────────────────────────────────
add_section_title(doc, 6, "Requerimientos Funcionales")
doc.add_paragraph()
add_body(doc, "A continuación se presentan los 40 requerimientos funcionales de VetCare, "
    "organizados en 6 módulos.", indent=False)

modulos = [
    ("MÓDULO 1 — Autenticación y Usuarios", [
        ("RF-01", "Iniciar sesión",     "El sistema permite autenticar al usuario con nombre de usuario y contraseña.", "Alta"),
        ("RF-02", "Cerrar sesión",      "El sistema cierra la sesión activa y regresa al menú de autenticación.", "Alta"),
        ("RF-03", "Cambiar contraseña", "El sistema permite modificar la contraseña ingresando la clave antigua y la nueva.", "Media"),
        ("RF-04", "Registrar usuario",  "El administrador crea nuevos usuarios con rol: admin o recepcionista.", "Media"),
    ]),
    ("MÓDULO 2 — Gestión de Mascotas", [
        ("RF-05", "Registrar mascota",     "El sistema registra una mascota con nombre, especie, raza, edad, sexo, peso y dueño.", "Alta"),
        ("RF-06", "Modificar mascota",     "El sistema permite actualizar los datos de una mascota existente.", "Alta"),
        ("RF-07", "Eliminar mascota",      "El sistema da de baja a una mascota previa confirmación del usuario.", "Media"),
        ("RF-08", "Buscar por nombre",     "El sistema busca mascotas ingresando el nombre o parte de él.", "Alta"),
        ("RF-09", "Buscar por dueño",      "El sistema filtra mascotas por nombre o DNI del dueño.", "Alta"),
        ("RF-10", "Listar mascotas",       "El sistema muestra el listado completo de mascotas con datos básicos.", "Alta"),
        ("RF-11", "Ver ficha completa",    "El sistema muestra todos los atributos de una mascota seleccionada.", "Media"),
        ("RF-12", "Registrar especie",     "El sistema clasifica la mascota por especie: Perro, Gato o Ave.", "Alta"),
        ("RF-13", "Registrar raza",        "El sistema permite ingresar la raza específica según la especie.", "Media"),
        ("RF-14", "Filtrar por especie",   "El sistema lista mascotas filtrando por especie seleccionada.", "Media"),
        ("RF-15", "Registrar observaciones","El sistema permite agregar notas sobre alergias o condiciones especiales.", "Alta"),
    ]),
    ("MÓDULO 3 — Gestión de Dueños", [
        ("RF-16", "Registrar dueño",       "El sistema registra un dueño con nombre, DNI, dirección, teléfono y correo.", "Alta"),
        ("RF-17", "Modificar dueño",       "El sistema permite actualizar la información de un dueño existente.", "Alta"),
        ("RF-18", "Eliminar dueño",        "El sistema elimina un dueño sin mascotas activas, previa confirmación.", "Baja"),
        ("RF-19", "Buscar por nombre",     "El sistema busca dueños ingresando nombre o parte de él.", "Alta"),
        ("RF-20", "Buscar por DNI",        "El sistema busca un dueño por su número de DNI.", "Alta"),
        ("RF-21", "Listar dueños",         "El sistema muestra la lista completa de dueños registrados.", "Media"),
        ("RF-22", "Ver mascotas del dueño","El sistema muestra todas las mascotas vinculadas a un dueño específico.", "Alta"),
    ]),
    ("MÓDULO 4 — Gestión de Citas", [
        ("RF-23", "Registrar cita",        "El sistema agenda una cita con mascota, veterinario, fecha, hora y motivo.", "Alta"),
        ("RF-24", "Modificar cita",        "El sistema cambia la fecha, hora o veterinario de una cita programada.", "Alta"),
        ("RF-25", "Cancelar cita",         "El sistema cancela una cita pendiente, registrando el motivo.", "Media"),
        ("RF-26", "Citas del día",         "El sistema muestra todas las citas programadas para la fecha actual.", "Alta"),
        ("RF-27", "Buscar por mascota",    "El sistema muestra el historial de citas de una mascota específica.", "Alta"),
        ("RF-28", "Buscar por fecha",      "El sistema consulta citas en una fecha específica.", "Media"),
        ("RF-29", "Listar citas pendientes","El sistema muestra todas las citas aún no atendidas.", "Alta"),
        ("RF-30", "Confirmar asistencia",  "El sistema marca una cita como 'Atendida' o 'No asistió'.", "Alta"),
        ("RF-31", "Motivo de la cita",     "El sistema permite ingresar el motivo de consulta al crear la cita.", "Alta"),
    ]),
    ("MÓDULO 5 — Gestión de Veterinarios", [
        ("RF-32", "Registrar veterinario", "El sistema registra un veterinario con nombre, especialidad y colegiatura.", "Alta"),
        ("RF-33", "Modificar veterinario", "El sistema actualiza la información de un veterinario registrado.", "Media"),
        ("RF-34", "Desactivar veterinario","El sistema desactiva un veterinario sin eliminarlo del registro.", "Baja"),
        ("RF-35", "Listar veterinarios",   "El sistema muestra la lista de todos los veterinarios activos.", "Alta"),
        ("RF-36", "Ver agenda",            "El sistema muestra todas las citas asignadas a un veterinario específico.", "Alta"),
    ]),
    ("MÓDULO 6 — Historial Médico y Reportes", [
        ("RF-37", "Registrar consulta",    "El sistema guarda el resultado de una cita: diagnóstico, observaciones y tratamiento.", "Alta"),
        ("RF-38", "Registrar tratamiento", "El sistema registra el medicamento o procedimiento indicado por el veterinario.", "Alta"),
        ("RF-39", "Ver historial",         "El sistema muestra el historial médico completo de una mascota seleccionada.", "Alta"),
        ("RF-40", "Reporte por período",   "El sistema genera un reporte de consultas realizadas en un rango de fechas.", "Media"),
    ]),
]

for mod_nombre, rfs in modulos:
    add_rf_table(doc, mod_nombre, rfs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── SECCIÓN 7: HISTORIAS DE USUARIO ──────────────────────────────────────────
add_section_title(doc, 7, "Historias de Usuario")
doc.add_paragraph()
add_body(doc, "Las historias de usuario describen las funcionalidades del sistema desde la perspectiva "
    "del usuario, siguiendo el formato: Como [rol], quiero [acción] para [beneficio].", indent=False)
doc.add_paragraph()

historias = [
    ("HU-01", "Iniciar sesión en el sistema",
     "Recepcionista o Administrador",
     "ingresar con mi usuario y contraseña al sistema VetCare",
     "acceder a las funciones según mi rol asignado",
     "RF-01, RF-04",
     "El sistema muestra el menú principal correspondiente al rol del usuario.",
     "Si las credenciales son incorrectas, muestra: 'Usuario o contraseña inválidos. Intente de nuevo.'"),
    ("HU-02", "Registrar una cita médica",
     "Recepcionista",
     "registrar una nueva cita seleccionando la mascota, veterinario, fecha y hora",
     "organizar la agenda de la clínica y evitar cruces de horario",
     "RF-23, RF-31",
     "La cita se registra y aparece en el listado del día correspondiente.",
     "Si el veterinario ya tiene cita en ese horario, el sistema advierte el conflicto."),
    ("HU-03", "Ver historial médico de una mascota",
     "Veterinario",
     "acceder al historial médico completo de una mascota",
     "revisar diagnósticos y tratamientos anteriores antes de atenderla",
     "RF-39, RF-11",
     "El sistema muestra todas las consultas previas con fecha, diagnóstico y tratamiento.",
     "Si no hay historial: muestra 'Sin consultas médicas registradas para esta mascota.'"),
    ("HU-04", "Registrar resultado de consulta",
     "Veterinario",
     "registrar el diagnóstico y tratamiento de una consulta",
     "mantener el historial médico actualizado de la mascota",
     "RF-37, RF-38",
     "La consulta se guarda y queda vinculada al historial de la mascota.",
     "Si no se ingresa diagnóstico, el sistema solicita el campo antes de guardar."),
    ("HU-05", "Ver citas del día",
     "Recepcionista",
     "ver todas las citas programadas para la fecha actual",
     "organizar la atención y preparar los expedientes del día",
     "RF-26",
     "Lista todas las citas con hora, mascota, dueño y veterinario asignado.",
     "Si no hay citas: muestra 'No hay citas programadas para hoy.'"),
    ("HU-06", "Registrar nueva mascota",
     "Recepcionista",
     "registrar una nueva mascota con todos sus datos",
     "crear su ficha antes de la primera consulta",
     "RF-05, RF-12, RF-13, RF-15",
     "La mascota se registra vinculada al dueño correspondiente.",
     "Si no se ingresa especie, el sistema muestra 'Campo obligatorio'."),
    ("HU-07", "Buscar dueño por DNI",
     "Recepcionista",
     "buscar un dueño ingresando su número de DNI",
     "verificar si ya está registrado y ver sus mascotas asociadas",
     "RF-20, RF-22",
     "El sistema muestra los datos del dueño y el listado de sus mascotas.",
     "Si el DNI no existe: 'Dueño no encontrado. ¿Desea registrarlo?'"),
    ("HU-08", "Cancelar una cita",
     "Recepcionista",
     "cancelar una cita pendiente registrando el motivo",
     "liberar el horario del veterinario cuando el dueño avisa que no asistirá",
     "RF-25",
     "La cita cambia a 'Cancelada' y el horario queda disponible.",
     "No se puede cancelar una cita con estado 'Atendida'."),
    ("HU-09", "Ver agenda del veterinario",
     "Administrador",
     "ver la agenda completa de un veterinario específico",
     "distribuir la carga de trabajo equitativamente entre el equipo médico",
     "RF-36",
     "Muestra todas las citas del veterinario con fecha, hora y mascota.",
     "Si no tiene citas: 'Este veterinario no tiene citas programadas.'"),
    ("HU-10", "Generar reporte de consultas",
     "Administrador",
     "generar un reporte de consultas realizadas en un período de fechas",
     "evaluar la productividad de la clínica y tomar decisiones de gestión",
     "RF-40",
     "Muestra número de consultas, mascotas atendidas y veterinarios del período.",
     "Si la fecha inicio es mayor a la fecha fin, el sistema muestra error de validación."),
]

for hu in historias:
    add_hu_card(doc, *hu)

doc.add_page_break()

# ── SECCIÓN 8: DIAGRAMA DE CLASES ─────────────────────────────────────────────
add_section_title(doc, 8, "Diagrama de Clases")
doc.add_paragraph()

add_subsection_title(doc, "8.1  Clases del Sistema")

clases_data = [
    ("Persona", "Abstracta", "-nombre:String\n-dni:String\n-telefono:String\n-correo:String",
     "+getNombre()\n+getDni()\n+getTelefono()", "Superclase de Dueno y Veterinario"),
    ("Dueno", "extends Persona", "id:int\n-direccion:String",
     "+registrar()\n+buscarPorDni()\n+listarMascotas()", "Hereda Persona · asociada a Mascota"),
    ("Veterinario", "extends Persona", "-especialidad:String\n-colegiatura:String\n-activo:boolean",
     "+registrar()\n+verAgenda()", "Hereda Persona · asociada a Cita"),
    ("Animal", "Abstracta", "-nombre:String\n-raza:String\n-edad:int\n-sexo:String\n-peso:double",
     "+getNombre()\n+getEspecie():String", "Superclase de Perro, Gato y Ave"),
    ("Perro", "extends Animal", "-tamano:String", "+getEspecie()→\"Perro\"", "Hereda Animal"),
    ("Gato",  "extends Animal", "-esCastrado:boolean", "+getEspecie()→\"Gato\"", "Hereda Animal"),
    ("Ave",   "extends Animal", "-tipoPico:String", "+getEspecie()→\"Ave\"", "Hereda Animal"),
    ("Mascota", "—", "-id:int\n-animal:Animal\n-dueno:Dueno\n-historial:ArrayList<ConsultaMedica>",
     "+registrar()\n+buscar()\n+verHistorial()", "Composición Animal(1..1) · ConsultaMedica(0..*)"),
    ("Cita", "—", "-id:int\n-mascota:Mascota\n-veterinario:Veterinario\n-fecha:String\n-hora:String\n-estado:String",
     "+registrar()\n+cancelar()\n+confirmar()", "Asociación Mascota · Veterinario"),
    ("ConsultaMedica", "—", "-id:int\n-fecha:String\n-diagnostico:String\n-tratamiento:String",
     "+registrarConsulta(d)\n+registrarConsulta(d,t)\n+registrarConsulta(d,t,o)", "Sobrecarga de métodos"),
    ("GestorMascotas", "—", "-mascotas:ArrayList<Mascota>",
     "+agregar()\n+buscarPorNombre()\n+filtrarPorEspecie()", "Composición Mascota(0..*)"),
    ("GestorCitas", "—", "-citas:ArrayList<Cita>",
     "+agendar()\n+verDelDia()\n+pendientes()", "Composición Cita(0..*)"),
    ("SistemaVetCare", "—", "-gestorMascotas\n-gestorCitas\n-vet:ArrayList<Veterinario>",
     "+iniciar()\n+menu()", "Orquesta GestorMascotas y GestorCitas"),
]

tbl_cls = doc.add_table(rows=1, cols=5)
tbl_cls.style = 'Table Grid'
for i, h in enumerate(['Clase', 'Tipo', 'Atributos', 'Métodos', 'Relación']):
    tbl_cls.rows[0].cells[i].text = h
    run_h = tbl_cls.rows[0].cells[i].paragraphs[0].runs[0]
    run_h.font.bold = True
    run_h.font.size = Pt(9)
    run_h.font.color.rgb = BLANCO
    set_cell_color(tbl_cls.rows[0].cells[i], '145A32')
    tbl_cls.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for idx, (clase, tipo, attrs, methods, rel) in enumerate(clases_data):
    row = tbl_cls.add_row()
    data = [clase, tipo, attrs, methods, rel]
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    for i, txt in enumerate(data):
        row.cells[i].text = txt
        if row.cells[i].paragraphs:
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = GRIS_TEXTO
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = VERDE_OSCURO
        set_cell_color(row.cells[i], bg)

for i, w in enumerate([Cm(2.5), Cm(2.5), Cm(4.0), Cm(4.0), Cm(3.0)]):
    for cell in tbl_cls.columns[i].cells:
        cell.width = w

add_subsection_title(doc, "8.2  Resumen de Relaciones")
relaciones = [
    "Herencia:      Persona → Dueno   |   Persona → Veterinario",
    "Herencia:      Animal  → Perro   |   Animal  → Gato   |   Animal → Ave",
    "Composición:   Mascota ◆─── Animal (1..1)  (el animal forma parte de la mascota)",
    "Composición:   Mascota ◆─── ConsultaMedica (0..*)  (historial de la mascota)",
    "Composición:   GestorMascotas ◆─── Mascota (0..*)",
    "Composición:   GestorCitas ◆─── Cita (0..*)",
    "Asociación:    Mascota ──── Dueno",
    "Asociación:    Cita ──── Mascota  |  Cita ──── Veterinario",
    "Sobrecarga:    ConsultaMedica.registrar() (3 versiones del método)",
]
for rel in relaciones:
    add_bullet(doc, rel)

add_body(doc, "Diagrama de clases UML generado para el sistema VetCare:")

img_uml = os.path.join(SCRIPT_DIR, "Semana3", "VetCare_Documentos", "uml_vetcare.png")
if os.path.exists(img_uml):
    p_uml = doc.add_paragraph()
    p_uml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_uml.add_run().add_picture(img_uml, width=Cm(15))
    p_caption_uml = doc.add_paragraph("Figura 2. Diagrama de Clases UML — VetCare (13 clases)")
    p_caption_uml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption_uml.runs[0].font.size = Pt(9)
    p_caption_uml.runs[0].font.italic = True
    p_caption_uml.runs[0].font.color.rgb = RGBColor(0x64, 0x64, 0x64)

doc.add_page_break()

# ── SECCIÓN 9: CRITERIOS DE ACEPTACIÓN ────────────────────────────────────────
add_section_title(doc, 9, "Criterios de Aceptación")
doc.add_paragraph()
add_body(doc, "Se definen los criterios de aceptación para los requerimientos funcionales de mayor "
    "prioridad del sistema VetCare.", indent=False)

criterios = [
    ("CA-01", "Autenticación de usuarios",
     "El sistema solo debe conceder acceso cuando el usuario y contraseña coincidan exactamente con los datos registrados. Ante tres intentos fallidos, el sistema debe bloquear el acceso por 5 minutos.",
     "RF-01"),
    ("CA-02", "Registro de mascota obligatorio",
     "No debe ser posible guardar una mascota sin nombre, especie y dueño asignado. El sistema debe validar estos campos antes de confirmar el registro.",
     "RF-05, RF-12"),
    ("CA-03", "Unicidad de DNI del dueño",
     "El sistema no debe permitir registrar dos dueños con el mismo número de DNI. Ante un DNI duplicado, debe mostrar el mensaje: 'Ya existe un dueño con ese DNI.'",
     "RF-16, RF-20"),
    ("CA-04", "Conflicto de horario en citas",
     "El sistema debe verificar si el veterinario seleccionado ya tiene una cita en la fecha y hora indicadas. Si hay conflicto, debe mostrar el aviso correspondiente antes de permitir el registro.",
     "RF-23"),
    ("CA-05", "Historial médico visible y ordenado",
     "El historial médico de una mascota debe mostrarse en orden cronológico descendente (más reciente primero), con fecha, diagnóstico y tratamiento en cada registro.",
     "RF-39"),
    ("CA-06", "Sobrecarga del método registrarConsulta()",
     "El sistema debe aceptar registrar una consulta con uno (diagnóstico), dos (diagnóstico + tratamiento) o tres parámetros (diagnóstico + tratamiento + observaciones), sin lanzar excepciones.",
     "RF-37"),
    ("CA-07", "Persistencia de datos entre sesiones",
     "Los datos registrados en una sesión deben estar disponibles al reiniciar el sistema, sin pérdida de información, gracias a la persistencia en archivos .txt.",
     "RF-37–RF-40"),
    ("CA-08", "Filtro por especie funcional",
     "Al filtrar mascotas por especie (Perro, Gato o Ave), el sistema debe mostrar únicamente las mascotas de esa categoría, sin incluir mascotas de otras especies.",
     "RF-14"),
    ("CA-09", "Reporte de consultas por período",
     "El reporte debe incluir el total de consultas, los nombres de las mascotas atendidas y los veterinarios que las atendieron, dentro del rango de fechas ingresado. Si la fecha de inicio supera a la fecha de fin, el sistema debe mostrar un mensaje de error.",
     "RF-40"),
]

for ca_id, nombre, criterio, rf_rel in criterios:
    tbl_ca = doc.add_table(rows=1, cols=4)
    tbl_ca.style = 'Table Grid'
    for i, (label, val) in enumerate([('ID', ca_id), ('Nombre', nombre), ('Criterio de aceptación', criterio), ('RF', rf_rel)]):
        cell = tbl_ca.rows[0].cells[i]
        if i == 0:
            p_h = cell.paragraphs[0]
            run_h = p_h.add_run(label + "\n")
            run_h.font.bold = True
            run_h.font.size = Pt(8.5)
            run_h.font.color.rgb = VERDE_OSCURO
            run_v = p_h.add_run(val)
            run_v.font.bold = True
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = VERDE_OSCURO
            set_cell_color(cell, 'D5F5E3')
        elif i == 1:
            p_h = cell.paragraphs[0]
            run_h = p_h.add_run(label + "\n")
            run_h.font.bold = True
            run_h.font.size = Pt(8.5)
            run_h.font.color.rgb = AMBAR
            run_v = p_h.add_run(val)
            run_v.font.bold = True
            run_v.font.size = Pt(9.5)
            run_v.font.color.rgb = GRIS_TEXTO
            set_cell_color(cell, 'FEF9E7')
        elif i == 2:
            p_h = cell.paragraphs[0]
            run_h = p_h.add_run(label + "\n")
            run_h.font.bold = True
            run_h.font.size = Pt(8.5)
            run_h.font.color.rgb = GRIS_TEXTO
            run_v = p_h.add_run(val)
            run_v.font.size = Pt(9.5)
            run_v.font.color.rgb = GRIS_TEXTO
            set_cell_color(cell, 'FFFFFF')
        else:
            p_h = cell.paragraphs[0]
            run_h = p_h.add_run(label + "\n")
            run_h.font.bold = True
            run_h.font.size = Pt(8.5)
            run_h.font.color.rgb = VERDE_OSCURO
            run_v = p_h.add_run(val)
            run_v.font.size = Pt(9)
            run_v.font.color.rgb = VERDE_MEDIO
            set_cell_color(cell, 'EAFAF1')

    for i, w in enumerate([Cm(1.3), Cm(3.2), Cm(10.0), Cm(1.5)]):
        tbl_ca.columns[i].cells[0].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── PIE DE PÁGINA ──────────────────────────────────────────────────────────────
doc.add_page_break()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_f = p_footer._p.get_or_add_pPr()
shd_f = OxmlElement('w:shd')
shd_f.set(qn('w:fill'), '145A32')
shd_f.set(qn('w:val'), 'clear')
pPr_f.append(shd_f)
r_f = p_footer.add_run("  VetCare — Sistema de Gestión de Clínica Veterinaria  ·  UPN 2026  "
    "·  Técnicas de Programación Orientada a Objetos  ")
r_f.font.size = Pt(9)
r_f.font.color.rgb = AMBAR_CLARO

# ── GUARDAR ────────────────────────────────────────────────────────────────────
out_path = os.path.join(SCRIPT_DIR, "Proyecto 2026", "Vetcare", "Documentacion", "VetCare_Avance2.docx")
doc.save(out_path)
print(f"Documento guardado en: {out_path}")
