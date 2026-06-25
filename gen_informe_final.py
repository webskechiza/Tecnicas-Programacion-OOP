"""
Script para generar VetCare_InformeFinal.docx
Informe completo (secciones 1-14) segun estructura oficial TPOO
Diseno: verde esmeralda + ambar dorado
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta de colores ──────────────────────────────────────────────────────────
VERDE_OSCURO = RGBColor(0x14, 0x5A, 0x32)
VERDE_MEDIO  = RGBColor(0x1E, 0x8B, 0x4F)
VERDE_CLARO  = RGBColor(0xD5, 0xF5, 0xE3)
AMBAR        = RGBColor(0xB7, 0x7A, 0x0C)
AMBAR_CLARO  = RGBColor(0xFD, 0xF2, 0xD0)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO   = RGBColor(0x2C, 0x3E, 0x50)
GRIS_CLARO   = RGBColor(0xF4, 0xF6, 0xF7)
ROJO         = RGBColor(0xC0, 0x39, 0x2B)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_DIR    = os.path.join(SCRIPT_DIR, "Proyecto 2026", "Vetcare", "Documentacion")

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_color(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_section_title(doc, numero, titulo):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '145A32')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    run_num = p.add_run(f"  {numero}.   ")
    run_num.font.size = Pt(13)
    run_num.font.bold = True
    run_num.font.color.rgb = AMBAR
    run_title = p.add_run(titulo.upper())
    run_title.font.size = Pt(13)
    run_title.font.bold = True
    run_title.font.color.rgb = BLANCO
    return p

def add_subsection_title(doc, titulo):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E8F8F0')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(f"  {titulo}")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = VERDE_OSCURO
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTO
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run_bullet = p.add_run("▸  ")
    run_bullet.font.color.rgb = VERDE_MEDIO
    run_bullet.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TEXTO
    return p

def add_bullet_red(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run_bullet = p.add_run("✗  ")
    run_bullet.font.color.rgb = ROJO
    run_bullet.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TEXTO
    return p

def add_alternativa(doc, restriccion, alternativa):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell1 = tbl.rows[0].cells[0]
    cell1.text = ""
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run("Restriccion:  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ROJO
    r2 = p1.add_run(restriccion)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRIS_TEXTO
    set_cell_color(cell1, 'FEF9E7')
    cell2 = tbl.rows[1].cells[0]
    cell2.text = ""
    p2 = cell2.paragraphs[0]
    r3 = p2.add_run("Alternativa de solucion:  ")
    r3.font.bold = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = VERDE_OSCURO
    r4 = p2.add_run(alternativa)
    r4.font.size = Pt(10)
    r4.font.color.rgb = GRIS_TEXTO
    set_cell_color(cell2, 'EAFAF1')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_rf_table(doc, modulo, rfs):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E8B4F')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(f"  {modulo}")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLANCO

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    for i, txt in enumerate(['ID', 'Nombre', 'Descripcion', 'Prioridad']):
        cell = tbl.rows[0].cells[i]
        cell.text = txt
        run_h = cell.paragraphs[0].runs[0]
        run_h.font.bold = True
        run_h.font.size = Pt(9.5)
        run_h.font.color.rgb = BLANCO
        set_cell_color(cell, '145A32')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (rid, nombre, desc, prio) in enumerate(rfs):
        row = tbl.add_row()
        bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
        for i, txt in enumerate([rid, nombre, desc, prio]):
            row.cells[i].text = txt
            run_d = row.cells[i].paragraphs[0].runs[0]
            run_d.font.size = Pt(9)
            run_d.font.color.rgb = GRIS_TEXTO
            set_cell_color(row.cells[i], bg)
        prio_cell = row.cells[3]
        prio_run = prio_cell.paragraphs[0].runs[0]
        if prio == 'Alta':
            prio_run.font.color.rgb = RGBColor(0x1A, 0x5E, 0x20)
            prio_run.font.bold = True
        elif prio == 'Media':
            prio_run.font.color.rgb = RGBColor(0x7D, 0x60, 0x08)
        elif prio == 'Baja':
            prio_run.font.color.rgb = RGBColor(0x78, 0x2F, 0x0E)
        prio_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, w in enumerate([Cm(1.5), Cm(3.5), Cm(9.0), Cm(2.0)]):
        for cell in tbl.columns[i].cells:
            cell.width = w

def add_hu_card(doc, hu_id, titulo, rol, quiero, para, rfs, escenario_ok, escenario_fail):
    p_title = doc.add_paragraph()
    pPr = p_title._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '145A32')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run_id = p_title.add_run(f"  {hu_id}  .  ")
    run_id.font.color.rgb = AMBAR
    run_id.font.bold = True
    run_id.font.size = Pt(10.5)
    run_t = p_title.add_run(titulo)
    run_t.font.color.rgb = BLANCO
    run_t.font.bold = True
    run_t.font.size = Pt(10.5)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    labels = ['Como', 'Quiero', 'Para', 'RF relacionados', 'Escenario exitoso']
    values = [rol, quiero, para, rfs, escenario_ok]
    for i, (label, value) in enumerate(zip(labels, values)):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = VERDE_OSCURO
        set_cell_color(tbl.rows[i].cells[0], 'E8F8F0')
        tbl.rows[i].cells[1].text = value
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = GRIS_TEXTO
        set_cell_color(tbl.rows[i].cells[1], 'FFFFFF')

    if escenario_fail:
        row_fail = tbl.add_row()
        row_fail.cells[0].text = 'Escenario fallido'
        row_fail.cells[0].paragraphs[0].runs[0].font.bold = True
        row_fail.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_fail.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAB, 0x27, 0x12)
        set_cell_color(row_fail.cells[0], 'FDEDEC')
        row_fail.cells[1].text = escenario_fail
        row_fail.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_color(row_fail.cells[1], 'FEF5F5')

    for cell in tbl.columns[0].cells:
        cell.width = Cm(3.5)
    for cell in tbl.columns[1].cells:
        cell.width = Cm(12.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCION DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ────────────────────────────────────────────────────────────────────
p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_inst._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '145A32'); shd.set(qn('w:val'), 'clear')
pPr.append(shd)
sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '160'); sp.set(qn('w:after'), '160')
pPr.append(sp)
r = p_inst.add_run("UNIVERSIDAD PRIVADA DEL NORTE")
r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = BLANCO

p_fac = doc.add_paragraph()
p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2 = p_fac._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd'); shd2.set(qn('w:fill'), '1E8B4F'); shd2.set(qn('w:val'), 'clear')
pPr2.append(shd2)
r1 = p_fac.add_run("Facultad de Ingenieria  |  Tecnicas de Programacion Orientada a Objetos")
r1.font.size = Pt(11); r1.font.color.rgb = BLANCO

doc.add_paragraph()

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr3 = p_name._p.get_or_add_pPr()
shd3 = OxmlElement('w:shd'); shd3.set(qn('w:fill'), 'EAFAF1'); shd3.set(qn('w:val'), 'clear')
pPr3.append(shd3)
sp3 = OxmlElement('w:spacing'); sp3.set(qn('w:before'), '240'); sp3.set(qn('w:after'), '80')
pPr3.append(sp3)
r_name = p_name.add_run("VETCARE")
r_name.font.size = Pt(36); r_name.font.bold = True; r_name.font.color.rgb = VERDE_OSCURO

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr4 = p_sub._p.get_or_add_pPr()
shd4 = OxmlElement('w:shd'); shd4.set(qn('w:fill'), 'EAFAF1'); shd4.set(qn('w:val'), 'clear')
pPr4.append(shd4)
r_sub = p_sub.add_run("Sistema de Gestion de Clinica Veterinaria")
r_sub.font.size = Pt(16); r_sub.font.color.rgb = AMBAR; r_sub.font.bold = True

p_tipo = doc.add_paragraph()
p_tipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr5 = p_tipo._p.get_or_add_pPr()
shd5 = OxmlElement('w:shd'); shd5.set(qn('w:fill'), 'EAFAF1'); shd5.set(qn('w:val'), 'clear')
pPr5.append(shd5)
sp5 = OxmlElement('w:spacing'); sp5.set(qn('w:before'), '80'); sp5.set(qn('w:after'), '240')
pPr5.append(sp5)
r_av = p_tipo.add_run("INFORME FINAL DE PROYECTO INTEGRADOR")
r_av.font.size = Pt(13); r_av.font.color.rgb = VERDE_MEDIO; r_av.font.bold = True

doc.add_paragraph()

tbl_info = doc.add_table(rows=5, cols=2)
tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER
datos = [
    ('Integrante:',  'Kevin Oswaldo Chirinos Zapata  -  N00521954'),
    ('Grupo:',       '28'),
    ('Docente:',     'Martin Eduardo Torres Rodriguez'),
    ('Curso:',       'Tecnicas de Programacion Orientada a Objetos'),
    ('Ciclo / Ano:', '2026-I'),
]
for i, (label, value) in enumerate(datos):
    tbl_info.rows[i].cells[0].text = label
    run_l = tbl_info.rows[i].cells[0].paragraphs[0].runs[0]
    run_l.font.bold = True; run_l.font.color.rgb = VERDE_OSCURO; run_l.font.size = Pt(10.5)
    set_cell_color(tbl_info.rows[i].cells[0], 'D5F5E3')
    tbl_info.rows[i].cells[1].text = value
    run_v = tbl_info.rows[i].cells[1].paragraphs[0].runs[0]
    run_v.font.size = Pt(10.5); run_v.font.color.rgb = GRIS_TEXTO
    set_cell_color(tbl_info.rows[i].cells[1], 'F8FDF9')
for cell in tbl_info.columns[0].cells:
    cell.width = Cm(4)
for cell in tbl_info.columns[1].cells:
    cell.width = Cm(10)

doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_l = p_line._p.get_or_add_pPr()
shd_l = OxmlElement('w:shd'); shd_l.set(qn('w:fill'), 'B77A0C'); shd_l.set(qn('w:val'), 'clear')
pPr_l.append(shd_l)
sp_l = OxmlElement('w:spacing'); sp_l.set(qn('w:before'), '0'); sp_l.set(qn('w:after'), '0')
pPr_l.append(sp_l)
p_line.add_run(" " * 80).font.size = Pt(4)

doc.add_page_break()

# ── INDICE ─────────────────────────────────────────────────────────────────────
p_idx = doc.add_paragraph()
p_idx.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_idx = p_idx._p.get_or_add_pPr()
shd_idx = OxmlElement('w:shd'); shd_idx.set(qn('w:fill'), '145A32'); shd_idx.set(qn('w:val'), 'clear')
pPr_idx.append(shd_idx)
r_idx = p_idx.add_run("  INDICE DE CONTENIDOS  ")
r_idx.font.size = Pt(14); r_idx.font.bold = True; r_idx.font.color.rgb = BLANCO

indice_items = [
    ("Introduccion",                                              "3"),
    ("1.   Definicion del Problema",                             "4"),
    ("2.   Antecedentes",                                        "5"),
    ("3.   Restricciones Realistas y Alternativas de Solucion",  "6"),
    ("4.   Objetivos del Proyecto",                              "8"),
    ("5.   Alcance de la Solucion",                              "9"),
    ("6.   Requerimientos Funcionales (40 RF)",                  "10"),
    ("7.   Historias de Usuario (10 HU)",                        "14"),
    ("8.   Diagrama de Clases UML",                              "19"),
    ("9.   Criterios de Aceptacion",                             "20"),
    ("10.  Implementacion del Proyecto",                         "22"),
    ("11.  Resultados",                                          "24"),
    ("12.  Conclusiones",                                        "26"),
    ("13.  Bibliografia",                                        "27"),
    ("14.  Anexos",                                              "28"),
]
for item, page in indice_items:
    p_i = doc.add_paragraph()
    p_i.paragraph_format.left_indent = Cm(1.0)
    p_i.paragraph_format.space_after = Pt(3)
    tab = p_i.add_run(item)
    tab.font.size = Pt(10.5); tab.font.color.rgb = GRIS_TEXTO
    dots = p_i.add_run("  " + "." * (62 - len(item)) + "  " + page)
    dots.font.size = Pt(10.5); dots.font.color.rgb = VERDE_MEDIO

doc.add_page_break()

# ── INTRODUCCION ───────────────────────────────────────────────────────────────
p_int_title = doc.add_paragraph()
pPr_int = p_int_title._p.get_or_add_pPr()
shd_int = OxmlElement('w:shd'); shd_int.set(qn('w:fill'), '145A32'); shd_int.set(qn('w:val'), 'clear')
pPr_int.append(shd_int)
r_int = p_int_title.add_run("  INTRODUCCION")
r_int.font.size = Pt(13); r_int.font.bold = True; r_int.font.color.rgb = BLANCO

add_body(doc, "VetCare es un sistema de gestion disenado para digitalizar y optimizar los procesos internos "
    "de las clinicas veterinarias de pequeno y mediano tamano. Este proyecto surge de la necesidad urgente "
    "de modernizar la gestion de pacientes animales, la programacion de citas medicas y el mantenimiento "
    "de historiales clinicos que, en la mayoria de los establecimientos veterinarios del Peru, aun se "
    "realizan de forma manual y desorganizada.", indent=False)

add_body(doc, "El sistema ha sido desarrollado en Java SE aplicando los principios de la Programacion "
    "Orientada a Objetos (POO). Este enfoque permite representar las entidades del negocio veterinario "
    "(mascotas, duenos, veterinarios y citas) como clases con atributos y metodos propios, garantizando "
    "un codigo modular, reutilizable y facil de mantener.", indent=False)

add_body(doc, "El presente documento constituye el informe final del proyecto integrador e incluye la "
    "definicion del problema, los antecedentes, las restricciones con sus alternativas de solucion, los "
    "objetivos, el alcance funcional, los 40 requerimientos funcionales del sistema, las 10 historias de "
    "usuario, el diagrama de clases UML, los criterios de aceptacion, la descripcion de la implementacion "
    "en Java, los resultados obtenidos, las conclusiones, la bibliografia y los anexos con el codigo fuente "
    "y el enlace al repositorio GitHub.", indent=False)

doc.add_page_break()

# ── SECCION 1: DEFINICION DEL PROBLEMA ────────────────────────────────────────
add_section_title(doc, 1, "Definicion del Problema")
doc.add_paragraph()

add_subsection_title(doc, "1.1  Situacion Problematica")
add_body(doc, "Las clinicas veterinarias de pequeno y mediano tamano en el Peru operan, en su mayoria, "
    "con procesos manuales que limitan su eficiencia operativa. El registro de mascotas, la programacion "
    "de citas y el mantenimiento del historial medico se realizan en cuadernos fisicos o planillas de "
    "Excel, lo que genera duplicidad de datos, perdida de informacion y demoras en la atencion.")

add_body(doc, "Esta situacion se agrava en periodos de alta demanda (campanas de vacunacion o desparasitacion "
    "masiva) donde la ausencia de un sistema organizado de citas provoca conflictos de horario y una atencion "
    "deficiente. Asimismo, la falta de un historial medico digital centralizado impide al veterinario "
    "acceder rapidamente al historial clinico del paciente durante la consulta.")

add_subsection_title(doc, "1.2  Problema Identificado")
p_q = doc.add_paragraph()
p_q.paragraph_format.left_indent = Cm(0.5)
pPr_q = p_q._p.get_or_add_pPr()
shd_q = OxmlElement('w:shd'); shd_q.set(qn('w:fill'), 'EAFAF1'); shd_q.set(qn('w:val'), 'clear')
pPr_q.append(shd_q)
r_q = p_q.add_run("De que manera el desarrollo de un sistema de gestion de clinica veterinaria en Java "
    "orientado a objetos puede digitalizar los procesos de registro de mascotas, programacion de citas e "
    "historial medico, mejorando la eficiencia operativa y la calidad de atencion de la clinica?")
r_q.font.size = Pt(10.5); r_q.font.italic = True; r_q.font.color.rgb = VERDE_OSCURO

add_subsection_title(doc, "1.3  Diagrama de Ishikawa")
add_body(doc, 'Problema central: "Ineficiencia en la gestion de atencion de pacientes en clinicas veterinarias de pequeno tamano"')

ishikawa_categorias = {
    "Personal:":         ["Falta de capacitacion tecnologica del personal",
                          "Resistencia al cambio y adopcion de herramientas digitales",
                          "Escaso personal para atencion y administracion simultanea"],
    "Metodo:":           ["Ausencia de protocolos de registro estandarizados",
                          "Programacion de citas sin sistema de confirmacion",
                          "Sin proceso de seguimiento post-consulta al dueno"],
    "Tecnologia:":       ["Ausencia de software especializado para veterinarias",
                          "Dependencia de cuadernos fisicos y hojas de calculo",
                          "Sin acceso a historial medico digital en tiempo real"],
    "Duenos:":           ["Sin recordatorios automaticos de citas o vacunas",
                          "Dificultad para consultar el historial de su mascota",
                          "Sin canal de comunicacion directo con la clinica"],
    "Historial medico:": ["Registros incompletos o ilegibles en papel",
                          "Informacion dispersa y sin centralizacion",
                          "Riesgo de perdida de datos ante extravios o siniestros"],
    "Finanzas:":         ["Cobros manuales sin registro sistematico",
                          "Sin reportes de ingresos por periodo",
                          "Dificultad para realizar cierres de caja"],
}
for categoria, items in ishikawa_categorias.items():
    add_subsection_title(doc, f"  {categoria}")
    for item in items:
        add_bullet(doc, item)

add_body(doc, "El diagrama completo de causa-efecto se presenta a continuacion:")

img_ishi = os.path.join(DOC_DIR, "ishikawa_vetcare.png")
if os.path.exists(img_ishi):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(img_ishi, width=Cm(14))
    p_cap = doc.add_paragraph("Figura 1. Diagrama de Ishikawa - VetCare")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.runs[0].font.size = Pt(9); p_cap.runs[0].font.italic = True
    p_cap.runs[0].font.color.rgb = RGBColor(0x64, 0x64, 0x64)

doc.add_page_break()

# ── SECCION 2: ANTECEDENTES ────────────────────────────────────────────────────
add_section_title(doc, 2, "Antecedentes")
doc.add_paragraph()
add_body(doc, "A continuacion se presentan tres antecedentes que respaldan el desarrollo de VetCare "
    "desde perspectivas internacional, nacional y tecnologica.", indent=False)

antecedentes = [
    ("Antecedente Internacional",
     "Barrientos, J., & Molina, P. (2020). Sistema de informacion para la gestion de clinicas veterinarias. "
     "Repositorio Universidad de Guayaquil, Ecuador.",
     "Este trabajo implemento un sistema de escritorio en Java para automatizar el registro de pacientes "
     "animales y la gestion de citas en clinicas veterinarias de Guayaquil, Ecuador. Los resultados evidenciaron "
     "una reduccion del 70% en el tiempo de registro de pacientes y una mejora significativa en la organizacion "
     "del historial medico. El proyecto guarda similitud directa con VetCare en cuanto al uso de Java y la "
     "gestion de citas como modulo central."),
    ("Antecedente Nacional",
     "Quispe, A., & Vargas, L. (2023). Desarrollo de software de escritorio para la gestion de consultas "
     "en clinicas veterinarias de Lima Metropolitana. Repositorio Institucional USIL, Lima, Peru.",
     "Este estudio desarrollo una aplicacion de escritorio para la clinica veterinaria 'Patitas Felices' "
     "en el distrito de San Borja, Lima. El sistema implementado redujo los errores en el registro de citas "
     "en un 85% y permitio centralizar el historial medico de mas de 500 mascotas. Los autores destacan que "
     "el uso de POO facilito la escalabilidad del sistema, alinandose con la propuesta arquitectonica de "
     "VetCare."),
    ("Antecedente Tecnologico",
     "Deitel, P. J., & Deitel, H. M. (2012). Java: Como programar (9.a ed.). Pearson Educacion.",
     "Esta referencia canonica del lenguaje Java fundamenta el paradigma de la Programacion Orientada a "
     "Objetos como base para el desarrollo de sistemas de gestion. Los autores demuestran que la representacion "
     "de entidades del mundo real (mascotas, medicos, citas) como clases con atributos y metodos propios "
     "constituye la estrategia optima para sistemas de escritorio escalables y mantenibles."),
]
for titulo, ref, desc in antecedentes:
    add_subsection_title(doc, titulo)
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.left_indent = Cm(0.5)
    run_ref = p_ref.add_run(ref)
    run_ref.font.size = Pt(10); run_ref.font.italic = True
    run_ref.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    add_body(doc, desc)

doc.add_page_break()

# ── SECCION 3: RESTRICCIONES ───────────────────────────────────────────────────
add_section_title(doc, 3, "Restricciones Realistas y Alternativas de Solucion")
doc.add_paragraph()
add_body(doc, "A continuacion se identifican las restricciones que afectan el desarrollo del proyecto VetCare, "
    "organizadas por categoria, junto con las alternativas de solucion propuestas para cada caso.", indent=False)

add_subsection_title(doc, "3.1  Restricciones Tecnicas")
rest_tecnicas = [
    ("El sistema se desarrolla en Java SE (JDK 17 o superior), sin uso de frameworks externos.",
     "Se aprovechan al maximo las librerias estandar de Java (java.util, java.io, java.time). "
     "En fases futuras se podria incorporar JDBC para acceso a base de datos sin cambiar el lenguaje."),
    ("La persistencia de datos se realiza mediante archivos de texto (.txt), sin base de datos relacional.",
     "Se implementa un gestor de archivos robusto con manejo de errores (try-catch) y separadores "
     "de campo definidos. En versiones posteriores se podria migrar a SQLite embebido sin cambiar la arquitectura."),
    ("La interfaz de usuario es de consola, sin interfaz grafica (GUI/Swing/JavaFX).",
     "Se disenan menus numerados intuitivos con validaciones en cada entrada y mensajes de error "
     "descriptivos, logrando una experiencia de usuario clara a pesar de la limitacion grafica."),
    ("El entorno de desarrollo es Eclipse IDE.",
     "El codigo es compatible con IntelliJ IDEA Community Edition y VS Code con extension Java."),
    ("El control de versiones se gestiona con Git y el repositorio se aloja en GitHub.",
     "En caso de inconvenientes de acceso, el repositorio puede duplicarse en GitLab como espejo de respaldo."),
]
for r, a in rest_tecnicas:
    add_alternativa(doc, r, a)

add_subsection_title(doc, "3.2  Restricciones Economicas")
rest_economicas = [
    ("Se utilizan exclusivamente herramientas gratuitas: Java SE, Eclipse IDE, Git y GitHub.",
     "Las versiones Community Edition de IntelliJ IDEA ofrecen funciones avanzadas de refactorizacion "
     "sin costo. Oracle Cloud Free Tier tambien puede usarse para pruebas de despliegue sin presupuesto."),
    ("No se cuenta con presupuesto para licencias de software, servidores ni servicios en la nube.",
     "GitHub provee repositorios publicos y GitHub Actions (CI) gratuitos, cubriendo las necesidades "
     "de control de versiones y automatizacion sin costo alguno."),
    ("El proyecto se desarrolla con los equipos personales de los integrantes del equipo.",
     "Los laboratorios de computo de la universidad estan disponibles para sesiones colaborativas "
     "presenciales cuando el equipo personal no este disponible."),
]
for r, a in rest_economicas:
    add_alternativa(doc, r, a)

add_subsection_title(doc, "3.3  Restricciones Operativas")
rest_operativas = [
    ("El tiempo de desarrollo esta limitado al ciclo academico 2026, con entregas parciales semanales.",
     "Se aplica una estrategia de priorizacion por modulos (MVP): primero mascotas, citas e historial; "
     "el modulo de reportes y autenticacion se integra en la fase final del ciclo."),
    ("Los integrantes tienen carga academica paralela, lo que limita la dedicacion exclusiva al proyecto.",
     "Se establecen reuniones semanales cortas de sincronizacion y se usa comunicacion asincrona "
     "para resolver bloqueos sin necesidad de reuniones adicionales."),
    ("El sistema no incluye integracion con sistemas externos (SUNAT, pasarelas de pago, APIs de terceros).",
     "El diseno de clases contempla interfaces preparadas para futura integracion REST, de modo que "
     "agregar conectores externos en versiones posteriores no requiera refactorizacion mayor."),
    ("El modulo de recordatorios automaticos por correo o SMS queda fuera del alcance del ciclo.",
     "Se implementa una funcion de 'agenda del dia' que el recepcionista consulta al iniciar la jornada."),
]
for r, a in rest_operativas:
    add_alternativa(doc, r, a)

doc.add_page_break()

# ── SECCION 4: OBJETIVOS ───────────────────────────────────────────────────────
add_section_title(doc, 4, "Objetivos del Proyecto")
doc.add_paragraph()

add_subsection_title(doc, "4.1  Objetivo General")
p_og = doc.add_paragraph()
p_og.paragraph_format.left_indent = Cm(0.5)
pPr_og = p_og._p.get_or_add_pPr()
shd_og = OxmlElement('w:shd'); shd_og.set(qn('w:fill'), 'EAFAF1'); shd_og.set(qn('w:val'), 'clear')
pPr_og.append(shd_og)
r_og = p_og.add_run('Desarrollar un sistema de gestion de clinica veterinaria denominado "VetCare" en Java '
    "orientado a objetos, que digitalice los procesos de registro de mascotas, gestion de citas medicas e "
    "historial clinico, mejorando la eficiencia operativa y la calidad de atencion de la clinica.")
r_og.font.size = Pt(11); r_og.font.color.rgb = GRIS_TEXTO

add_subsection_title(doc, "4.2  Objetivos Especificos")
oe_table = doc.add_table(rows=1, cols=2)
oe_table.style = 'Table Grid'
for hdr_txt, cell_obj in zip(['ID', 'Objetivo Especifico'], oe_table.rows[0].cells):
    cell_obj.text = hdr_txt
    cell_obj.paragraphs[0].runs[0].font.bold = True
    cell_obj.paragraphs[0].runs[0].font.size = Pt(10)
    cell_obj.paragraphs[0].runs[0].font.color.rgb = BLANCO
    set_cell_color(cell_obj, '145A32')
    cell_obj.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

oes = [
    ("OE-01", "Implementar un modulo de gestion de mascotas que permita registrar, consultar y mantener "
              "el historial medico de cada paciente animal."),
    ("OE-02", "Desarrollar un modulo de gestion de citas que organice la programacion de atenciones "
              "veterinarias con control de conflictos de horario."),
    ("OE-03", "Crear un modulo de gestion de clientes (duenos) que centralice la informacion de contacto "
              "y su vinculacion con las mascotas registradas."),
    ("OE-04", "Disenar un diagrama de clases UML que aplique herencia (Animal: Perro, Gato, Ave), "
              "composicion y asociacion entre las entidades del sistema."),
    ("OE-05", "Aplicar principios de POO (herencia, encapsulamiento, sobrecarga de metodos) y colecciones "
              "ArrayList para estructurar el codigo de forma modular y escalable."),
    ("OE-06", "Implementar persistencia de datos mediante archivos .txt con lectura y escritura en "
              "formato UTF-8 para conservar la informacion entre sesiones del sistema."),
]
for idx, (oe_id, oe_desc) in enumerate(oes):
    row = oe_table.add_row()
    row.cells[0].text = oe_id
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = VERDE_OSCURO
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    set_cell_color(row.cells[0], bg)
    row.cells[1].text = oe_desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = GRIS_TEXTO
    set_cell_color(row.cells[1], bg)

for cell in oe_table.columns[0].cells:
    cell.width = Cm(2.0)
for cell in oe_table.columns[1].cells:
    cell.width = Cm(14.0)

doc.add_page_break()

# ── SECCION 5: ALCANCE ─────────────────────────────────────────────────────────
add_section_title(doc, 5, "Alcance de la Solucion")
doc.add_paragraph()
add_body(doc, "El alcance define con precision que funcionalidades incluye VetCare en el ciclo 2026 "
    "y cuales quedan fuera del desarrollo actual.", indent=False)

add_subsection_title(doc, "Funcionalidades incluidas")
incluidas = [
    "Modulo de autenticacion con login (usuario admin / recepcionista).",
    "Gestion completa de mascotas: registro, eliminacion, busqueda por nombre y filtrado por especie.",
    "Clasificacion de mascotas por especie con herencia UML (Perro, Gato, Ave).",
    "Gestion de duenos: registro, listado, busqueda por DNI e historial de mascotas asociadas.",
    "Gestion de medicos veterinarios: registro, listado, agenda de citas y desactivacion.",
    "Gestion de citas: agendado con deteccion de conflictos, cancelacion y confirmacion de asistencia.",
    "Registro de consultas medicas: diagnostico, tratamiento y observaciones (sobrecarga de constructores).",
    "Historial medico digital por mascota.",
    "Persistencia de datos en archivos de texto (.txt) con carga al iniciar el sistema.",
    "Menu de consola con navegacion numerada e validaciones de entrada.",
]
for item in incluidas:
    add_bullet(doc, item)

add_subsection_title(doc, "Funcionalidades NO incluidas")
no_incluidas = [
    "Interfaz grafica de usuario (GUI / Swing / JavaFX).",
    "Conexion a base de datos relacional (MySQL, PostgreSQL).",
    "Modulo de facturacion, pagos o integracion con POS.",
    "Envio de recordatorios automaticos por correo electronico o SMS.",
    "Aplicacion web o movil.",
    "Modulo de farmacia o control de medicamentos.",
    "Integracion con sistemas externos de identificacion de mascotas.",
]
for item in no_incluidas:
    add_bullet_red(doc, item)

doc.add_page_break()

# ── SECCION 6: REQUERIMIENTOS FUNCIONALES ─────────────────────────────────────
add_section_title(doc, 6, "Requerimientos Funcionales")
doc.add_paragraph()
add_body(doc, "A continuacion se presentan los 40 requerimientos funcionales de VetCare, "
    "organizados en 6 modulos tematicos.", indent=False)

modulos = [
    ("MODULO 1 - Autenticacion y Usuarios", [
        ("RF-01", "Iniciar sesion",     "El sistema permite autenticar al usuario con nombre de usuario y contrasena.", "Alta"),
        ("RF-02", "Cerrar sesion",      "El sistema cierra la sesion activa y regresa al menu de autenticacion.", "Alta"),
        ("RF-03", "Cambiar contrasena", "El sistema permite modificar la contrasena ingresando la clave antigua y la nueva.", "Media"),
        ("RF-04", "Registrar usuario",  "El administrador crea nuevos usuarios con rol: admin o recepcionista.", "Media"),
    ]),
    ("MODULO 2 - Gestion de Mascotas", [
        ("RF-05", "Registrar mascota",      "El sistema registra una mascota con nombre, especie, raza, edad, sexo, peso y dueno.", "Alta"),
        ("RF-06", "Modificar mascota",      "El sistema permite actualizar los datos de una mascota existente.", "Alta"),
        ("RF-07", "Eliminar mascota",       "El sistema da de baja a una mascota previa confirmacion del usuario.", "Media"),
        ("RF-08", "Buscar por nombre",      "El sistema busca mascotas ingresando el nombre o parte de el.", "Alta"),
        ("RF-09", "Buscar por dueno",       "El sistema filtra mascotas por nombre o DNI del dueno.", "Alta"),
        ("RF-10", "Listar mascotas",        "El sistema muestra el listado completo de mascotas con datos basicos.", "Alta"),
        ("RF-11", "Ver ficha completa",     "El sistema muestra todos los atributos de una mascota seleccionada.", "Media"),
        ("RF-12", "Registrar especie",      "El sistema clasifica la mascota por especie: Perro, Gato o Ave.", "Alta"),
        ("RF-13", "Registrar raza",         "El sistema permite ingresar la raza especifica segun la especie.", "Media"),
        ("RF-14", "Filtrar por especie",    "El sistema lista mascotas filtrando por especie seleccionada.", "Media"),
        ("RF-15", "Registrar observaciones","El sistema permite agregar notas sobre alergias o condiciones especiales.", "Alta"),
    ]),
    ("MODULO 3 - Gestion de Duenos", [
        ("RF-16", "Registrar dueno",        "El sistema registra un dueno con nombre, DNI, direccion, telefono y correo.", "Alta"),
        ("RF-17", "Modificar dueno",        "El sistema permite actualizar la informacion de un dueno existente.", "Alta"),
        ("RF-18", "Eliminar dueno",         "El sistema elimina un dueno sin mascotas activas, previa confirmacion.", "Baja"),
        ("RF-19", "Buscar por nombre",      "El sistema busca duenos ingresando nombre o parte de el.", "Alta"),
        ("RF-20", "Buscar por DNI",         "El sistema busca un dueno por su numero de DNI.", "Alta"),
        ("RF-21", "Listar duenos",          "El sistema muestra la lista completa de duenos registrados.", "Media"),
        ("RF-22", "Ver mascotas del dueno", "El sistema muestra todas las mascotas vinculadas a un dueno especifico.", "Alta"),
    ]),
    ("MODULO 4 - Gestion de Citas", [
        ("RF-23", "Registrar cita",         "El sistema agenda una cita con mascota, veterinario, fecha, hora y motivo.", "Alta"),
        ("RF-24", "Modificar cita",         "El sistema cambia la fecha, hora o veterinario de una cita programada.", "Alta"),
        ("RF-25", "Cancelar cita",          "El sistema cancela una cita pendiente, registrando el motivo.", "Media"),
        ("RF-26", "Citas del dia",          "El sistema muestra todas las citas programadas para la fecha actual.", "Alta"),
        ("RF-27", "Buscar por mascota",     "El sistema muestra el historial de citas de una mascota especifica.", "Alta"),
        ("RF-28", "Buscar por fecha",       "El sistema consulta citas en una fecha especifica.", "Media"),
        ("RF-29", "Listar citas pendientes","El sistema muestra todas las citas aun no atendidas.", "Alta"),
        ("RF-30", "Confirmar asistencia",   "El sistema marca una cita como 'Atendida' o 'No asistio'.", "Alta"),
        ("RF-31", "Motivo de la cita",      "El sistema permite ingresar el motivo de consulta al crear la cita.", "Alta"),
    ]),
    ("MODULO 5 - Gestion de Veterinarios", [
        ("RF-32", "Registrar veterinario",  "El sistema registra un veterinario con nombre, especialidad y colegiatura.", "Alta"),
        ("RF-33", "Modificar veterinario",  "El sistema actualiza la informacion de un veterinario registrado.", "Media"),
        ("RF-34", "Desactivar veterinario", "El sistema desactiva un veterinario sin eliminarlo del registro.", "Baja"),
        ("RF-35", "Listar veterinarios",    "El sistema muestra la lista de todos los veterinarios activos.", "Alta"),
        ("RF-36", "Ver agenda",             "El sistema muestra todas las citas asignadas a un veterinario especifico.", "Alta"),
    ]),
    ("MODULO 6 - Historial Medico y Reportes", [
        ("RF-37", "Registrar consulta",     "El sistema guarda el resultado de una cita: diagnostico, observaciones y tratamiento.", "Alta"),
        ("RF-38", "Registrar tratamiento",  "El sistema registra el medicamento o procedimiento indicado por el veterinario.", "Alta"),
        ("RF-39", "Ver historial",          "El sistema muestra el historial medico completo de una mascota seleccionada.", "Alta"),
        ("RF-40", "Reporte por periodo",    "El sistema genera un reporte de consultas realizadas en un rango de fechas.", "Media"),
    ]),
]

for mod_nombre, rfs in modulos:
    add_rf_table(doc, mod_nombre, rfs)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── SECCION 7: HISTORIAS DE USUARIO ───────────────────────────────────────────
add_section_title(doc, 7, "Historias de Usuario")
doc.add_paragraph()
add_body(doc, "Las historias de usuario describen las funcionalidades del sistema desde la perspectiva "
    "del usuario, siguiendo el formato: Como [rol], quiero [accion] para [beneficio].", indent=False)
doc.add_paragraph()

historias = [
    ("HU-01", "Iniciar sesion en el sistema",
     "Recepcionista o Administrador",
     "ingresar con mi usuario y contrasena al sistema VetCare",
     "acceder a las funciones segun mi rol asignado",
     "RF-01, RF-04",
     "El sistema muestra el menu principal al validar las credenciales correctas.",
     "Si las credenciales son incorrectas, muestra: 'Usuario o contrasena invalidos. Intente de nuevo.' Tras 3 intentos fallidos, bloquea el acceso."),
    ("HU-02", "Registrar una cita medica",
     "Recepcionista",
     "registrar una nueva cita seleccionando la mascota, veterinario, fecha y hora",
     "organizar la agenda de la clinica y evitar cruces de horario",
     "RF-23, RF-31",
     "La cita se registra y aparece en el listado del dia correspondiente.",
     "Si el veterinario ya tiene cita en ese horario, el sistema muestra un mensaje de conflicto y no registra la cita."),
    ("HU-03", "Ver historial medico de una mascota",
     "Veterinario",
     "acceder al historial medico completo de una mascota",
     "revisar diagnosticos y tratamientos anteriores antes de atenderla",
     "RF-39, RF-11",
     "El sistema muestra todas las consultas previas con fecha, diagnostico y tratamiento.",
     "Si no hay historial: muestra 'Sin consultas medicas registradas para esta mascota.'"),
    ("HU-04", "Registrar resultado de consulta",
     "Veterinario",
     "registrar el diagnostico y tratamiento de una consulta",
     "mantener el historial medico actualizado de la mascota",
     "RF-37, RF-38",
     "La consulta se guarda y queda vinculada al historial de la mascota.",
     "Si no se ingresa diagnostico, el sistema solicita el campo antes de guardar."),
    ("HU-05", "Ver citas del dia",
     "Recepcionista",
     "ver todas las citas programadas para la fecha actual",
     "organizar la atencion y preparar los expedientes del dia",
     "RF-26",
     "Lista todas las citas con hora, mascota, dueno y veterinario asignado.",
     "Si no hay citas: muestra 'No hay citas programadas para hoy.'"),
    ("HU-06", "Registrar nueva mascota",
     "Recepcionista",
     "registrar una nueva mascota con todos sus datos",
     "crear su ficha antes de la primera consulta",
     "RF-05, RF-12, RF-13, RF-15",
     "La mascota se registra vinculada al dueno correspondiente.",
     "Si no se ingresa especie valida, el sistema muestra 'Especie no valida' y no guarda el registro."),
    ("HU-07", "Buscar dueno por DNI",
     "Recepcionista",
     "buscar un dueno ingresando su numero de DNI",
     "verificar si ya esta registrado y ver sus mascotas asociadas",
     "RF-20, RF-22",
     "El sistema muestra los datos del dueno y el listado de sus mascotas.",
     "Si el DNI no existe: 'Dueno no encontrado. Desea registrarlo?'"),
    ("HU-08", "Cancelar una cita",
     "Recepcionista",
     "cancelar una cita pendiente registrando el motivo",
     "liberar el horario del veterinario cuando el dueno avisa que no asistira",
     "RF-25",
     "La cita cambia a estado 'Cancelada' y el horario queda disponible para otra cita.",
     "No se puede cancelar una cita con estado 'Atendida'; el sistema muestra error de estado."),
    ("HU-09", "Ver agenda del veterinario",
     "Administrador",
     "ver la agenda completa de un veterinario especifico",
     "distribuir la carga de trabajo equitativamente entre el equipo medico",
     "RF-36",
     "Muestra todas las citas del veterinario con fecha, hora y mascota.",
     "Si no tiene citas: 'Este veterinario no tiene citas programadas.'"),
    ("HU-10", "Confirmar asistencia de una cita",
     "Recepcionista",
     "marcar si el paciente asistio o no a la cita programada",
     "mantener actualizado el estado de las citas y el historial de atencion",
     "RF-30",
     "La cita se actualiza a 'Atendida' o 'No asistio' segun la respuesta del recepcionista.",
     "No se puede confirmar una cita con estado 'Cancelada'; el sistema muestra error de estado."),
]

for hu in historias:
    add_hu_card(doc, *hu)

doc.add_page_break()

# ── SECCION 8: DIAGRAMA DE CLASES ─────────────────────────────────────────────
add_section_title(doc, 8, "Diagrama de Clases")
doc.add_paragraph()

add_subsection_title(doc, "8.1  Clases del Sistema")
clases_data = [
    ("Persona",         "Abstracta",        "-nombre:String\n-dni:String\n-telefono:String\n-correo:String",
     "+getNombre()\n+getDni()\n+getTelefono()", "Superclase de Dueno y Veterinario"),
    ("Dueno",           "extends Persona",  "-id:int\n-direccion:String",
     "+registrar()\n+buscarPorDni()\n+listarMascotas()", "Hereda Persona; asociada a Mascota"),
    ("Veterinario",     "extends Persona",  "-especialidad:String\n-colegiatura:String\n-activo:boolean",
     "+registrar()\n+verAgenda()\n+desactivar()", "Hereda Persona; asociada a Cita"),
    ("Animal",          "Abstracta",        "-nombre:String\n-raza:String\n-edad:int\n-sexo:String\n-peso:double",
     "+getNombre()\n+getEspecie():String", "Superclase de Perro, Gato y Ave"),
    ("Perro",           "extends Animal",   "-tamano:String",                "+getEspecie()->Perro", "Hereda Animal"),
    ("Gato",            "extends Animal",   "-esCastrado:boolean",           "+getEspecie()->Gato",  "Hereda Animal"),
    ("Ave",             "extends Animal",   "-tipoPico:String",              "+getEspecie()->Ave",   "Hereda Animal"),
    ("Mascota",         "Clase normal",     "-id:int\n-animal:Animal\n-dueno:Dueno\n-historial:ArrayList<ConsultaMedica>",
     "+registrar()\n+buscar()\n+agregarConsulta()\n+mostrarHistorial()", "Composicion: Animal(1), ConsultaMedica(0..*)"),
    ("Cita",            "Clase normal",     "-id:int\n-mascota:Mascota\n-veterinario:Veterinario\n-fecha:String\n-hora:String\n-estado:Estado",
     "+confirmarAsistencia()\n+cancelar()\n+marcarNoAsistio()", "Asociacion: Mascota y Veterinario"),
    ("ConsultaMedica",  "Clase normal",     "-id:int\n-fecha:String\n-diagnostico:String\n-tratamiento:String\n-observaciones:String",
     "+ConsultaMedica(id,fecha,diag)\n+ConsultaMedica(id,fecha,diag,trat)\n+ConsultaMedica(id,fecha,diag,trat,obs)", "Sobrecarga de 3 constructores"),
    ("GestorMascotas",  "Clase normal",     "-mascotas:ArrayList<Mascota>",
     "+agregar()\n+buscarPorNombre()\n+filtrarPorEspecie()\n+eliminar()", "Implementa IGestionable<Mascota>"),
    ("GestorCitas",     "Clase normal",     "-citas:ArrayList<Cita>",
     "+agregar()\n+pendientes()\n+porFecha()\n+citasDelDia()", "Implementa IGestionable<Cita>"),
    ("IGestionable<T>", "Interfaz",         "--",
     "+agregar(T)\n+buscarPorId(int)\n+listar()\n+eliminar(int)", "Implementada por GestorMascotas y GestorCitas"),
    ("SistemaVetCare",  "Clase normal",     "-gestorMascotas\n-gestorCitas\n-duenos:ArrayList<Dueno>\n-veterinarios:ArrayList<Veterinario>",
     "+iniciar()\n+autenticar()\n+menuPrincipal()", "Orquesta todos los gestores; contiene el menu principal"),
    ("PersistenciaArchivos","Clase normal", "-carpetaData:String",
     "+guardarDuenos()\n+guardarVeterinarios()\n+guardarMascotas()\n+cargarDuenos()\n+cargarVeterinarios()", "Lectura y escritura de archivos .txt UTF-8"),
]

tbl_cls = doc.add_table(rows=1, cols=5)
tbl_cls.style = 'Table Grid'
for i, h in enumerate(['Clase', 'Tipo', 'Atributos', 'Metodos', 'Relacion']):
    tbl_cls.rows[0].cells[i].text = h
    run_h = tbl_cls.rows[0].cells[i].paragraphs[0].runs[0]
    run_h.font.bold = True; run_h.font.size = Pt(9); run_h.font.color.rgb = BLANCO
    set_cell_color(tbl_cls.rows[0].cells[i], '145A32')
    tbl_cls.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for idx, (clase, tipo, attrs, methods, rel) in enumerate(clases_data):
    row = tbl_cls.add_row()
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    for i, txt in enumerate([clase, tipo, attrs, methods, rel]):
        row.cells[i].text = txt
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8.5); run.font.color.rgb = GRIS_TEXTO
                if i == 0:
                    run.font.bold = True; run.font.color.rgb = VERDE_OSCURO
        set_cell_color(row.cells[i], bg)

for i, w in enumerate([Cm(2.5), Cm(2.3), Cm(4.0), Cm(4.0), Cm(3.2)]):
    for cell in tbl_cls.columns[i].cells:
        cell.width = w

add_subsection_title(doc, "8.2  Resumen de Relaciones UML")
relaciones = [
    "Herencia:     Persona -> Dueno   |   Persona -> Veterinario",
    "Herencia:     Animal  -> Perro   |   Animal  -> Gato   |   Animal -> Ave",
    "Composicion:  Mascota contiene Animal (1..1) y ConsultaMedica (0..*)",
    "Composicion:  GestorMascotas contiene Mascota (0..*)",
    "Composicion:  GestorCitas contiene Cita (0..*)",
    "Asociacion:   Mascota -> Dueno (el dueno posee la mascota)",
    "Asociacion:   Cita -> Mascota  |  Cita -> Veterinario",
    "Interfaz:     IGestionable<T> implementada por GestorMascotas y GestorCitas",
    "Sobrecarga:   ConsultaMedica tiene 3 constructores (1, 2 y 3 parametros)",
]
for rel in relaciones:
    add_bullet(doc, rel)

img_uml = os.path.join(DOC_DIR, "uml_vetcare.png")
if os.path.exists(img_uml):
    add_body(doc, "Diagrama de clases UML generado para el sistema VetCare:")
    p_uml = doc.add_paragraph()
    p_uml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_uml.add_run().add_picture(img_uml, width=Cm(15))
    p_cap_uml = doc.add_paragraph("Figura 2. Diagrama de Clases UML - VetCare (15 clases)")
    p_cap_uml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_uml.runs[0].font.size = Pt(9); p_cap_uml.runs[0].font.italic = True
    p_cap_uml.runs[0].font.color.rgb = RGBColor(0x64, 0x64, 0x64)

doc.add_page_break()

# ── SECCION 9: CRITERIOS DE ACEPTACION ────────────────────────────────────────
add_section_title(doc, 9, "Criterios de Aceptacion")
doc.add_paragraph()
add_body(doc, "Se definen los criterios de aceptacion para los requerimientos funcionales de mayor "
    "prioridad del sistema VetCare.", indent=False)

criterios = [
    ("CA-01", "Autenticacion de usuarios",
     "El sistema solo concede acceso cuando el usuario y la contrasena coincidan exactamente con los datos registrados. Ante tres intentos fallidos, bloquea el acceso.",
     "RF-01"),
    ("CA-02", "Registro de mascota obligatorio",
     "No debe ser posible guardar una mascota sin nombre, especie y dueno asignado. El sistema valida estos campos antes de confirmar el registro.",
     "RF-05, RF-12"),
    ("CA-03", "Unicidad de DNI del dueno",
     "El sistema no permite registrar dos duenos con el mismo numero de DNI. Ante un DNI duplicado, muestra: 'Ya existe un dueno con ese DNI.'",
     "RF-16, RF-20"),
    ("CA-04", "Conflicto de horario en citas",
     "El sistema verifica si el veterinario seleccionado ya tiene una cita en la fecha y hora indicadas. Si hay conflicto, muestra el aviso correspondiente antes de permitir el registro.",
     "RF-23"),
    ("CA-05", "Historial medico visible y ordenado",
     "El historial medico de una mascota se muestra con fecha, diagnostico y tratamiento en cada registro, vinculado unicamente a la mascota seleccionada.",
     "RF-39"),
    ("CA-06", "Sobrecarga de constructores en ConsultaMedica",
     "El sistema acepta registrar una consulta con uno (diagnostico), dos (diagnostico + tratamiento) o tres parametros (diagnostico + tratamiento + observaciones) sin lanzar excepciones.",
     "RF-37"),
    ("CA-07", "Persistencia de datos entre sesiones",
     "Los datos registrados en una sesion deben estar disponibles al reiniciar el sistema, sin perdida de informacion, gracias a la persistencia en archivos .txt.",
     "RF-37, RF-38, RF-39, RF-40"),
    ("CA-08", "Filtro por especie funcional",
     "Al filtrar mascotas por especie (Perro, Gato o Ave), el sistema muestra unicamente las mascotas de esa categoria, sin incluir mascotas de otras especies.",
     "RF-14"),
    ("CA-09", "Cancelacion de cita con motivo",
     "Solo se pueden cancelar citas con estado 'Pendiente'. El sistema rechaza la cancelacion de citas 'Atendidas' y muestra un mensaje de error claro.",
     "RF-25"),
    ("CA-10", "Desactivacion de veterinario sin eliminacion",
     "Al desactivar un veterinario, este deja de aparecer en el listado de activos pero sus citas y datos historicos permanecen en el sistema.",
     "RF-34"),
]

for ca_id, nombre, criterio, rf_rel in criterios:
    tbl_ca = doc.add_table(rows=1, cols=4)
    tbl_ca.style = 'Table Grid'
    for i, (label, val) in enumerate([('ID', ca_id), ('Nombre', nombre), ('Criterio de aceptacion', criterio), ('RF', rf_rel)]):
        cell = tbl_ca.rows[0].cells[i]
        p_h = cell.paragraphs[0]
        run_h = p_h.add_run(label + "\n")
        run_h.font.bold = True; run_h.font.size = Pt(8.5)
        run_v = p_h.add_run(val)
        run_v.font.size = Pt(9.5) if i != 2 else Pt(9)
        if i == 0:
            run_h.font.color.rgb = VERDE_OSCURO; run_v.font.bold = True; run_v.font.color.rgb = VERDE_OSCURO
            set_cell_color(cell, 'D5F5E3')
        elif i == 1:
            run_h.font.color.rgb = AMBAR; run_v.font.bold = True; run_v.font.color.rgb = GRIS_TEXTO
            set_cell_color(cell, 'FEF9E7')
        elif i == 2:
            run_h.font.color.rgb = GRIS_TEXTO; run_v.font.color.rgb = GRIS_TEXTO
            set_cell_color(cell, 'FFFFFF')
        else:
            run_h.font.color.rgb = VERDE_OSCURO; run_v.font.color.rgb = VERDE_MEDIO
            set_cell_color(cell, 'EAFAF1')

    for i, w in enumerate([Cm(1.3), Cm(3.2), Cm(10.0), Cm(1.5)]):
        tbl_ca.columns[i].cells[0].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ── SECCION 10: IMPLEMENTACION DEL PROYECTO ───────────────────────────────────
add_section_title(doc, 10, "Implementacion del Proyecto")
doc.add_paragraph()
add_body(doc, "Esta seccion describe las decisiones tecnicas, herramientas y estrategias utilizadas "
    "en el desarrollo del sistema VetCare en Java.", indent=False)

add_subsection_title(doc, "10.1  Tecnologias Utilizadas")
tbl_tec = doc.add_table(rows=1, cols=3)
tbl_tec.style = 'Table Grid'
for i, h in enumerate(['Tecnologia', 'Version / Detalle', 'Uso en el proyecto']):
    tbl_tec.rows[0].cells[i].text = h
    run_h = tbl_tec.rows[0].cells[i].paragraphs[0].runs[0]
    run_h.font.bold = True; run_h.font.size = Pt(9.5); run_h.font.color.rgb = BLANCO
    set_cell_color(tbl_tec.rows[0].cells[i], '145A32')
    tbl_tec.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

tecnologias = [
    ("Java SE",        "JDK 17 LTS",         "Lenguaje de programacion principal. Aplicacion de POO, colecciones y manejo de archivos."),
    ("Eclipse IDE",    "2023-12 (4.30)",      "Entorno de desarrollo integrado. Compilacion, depuracion y gestion del workspace."),
    ("Git",            "2.44.x",             "Control de versiones distribuido. Registro de commits y ramas del proyecto."),
    ("GitHub",         "webskechiza/repo",   "Repositorio remoto. Respaldo y entrega del codigo fuente con historial de commits."),
    ("java.io / java.util", "Java SE stdlib","Persistencia en archivos .txt, colecciones ArrayList y Scanner para entrada del usuario."),
]
for idx, (tec, ver, uso) in enumerate(tecnologias):
    row = tbl_tec.add_row()
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    for i, txt in enumerate([tec, ver, uso]):
        row.cells[i].text = txt
        run_d = row.cells[i].paragraphs[0].runs[0]
        run_d.font.size = Pt(9.5); run_d.font.color.rgb = GRIS_TEXTO
        if i == 0:
            run_d.font.bold = True; run_d.font.color.rgb = VERDE_OSCURO
        set_cell_color(row.cells[i], bg)
for i, w in enumerate([Cm(3.5), Cm(3.5), Cm(9.0)]):
    for cell in tbl_tec.columns[i].cells:
        cell.width = w

doc.add_paragraph()
add_subsection_title(doc, "10.2  Manejo de Archivos en Java")
add_body(doc, "VetCare implementa persistencia de datos mediante la clase PersistenciaArchivos.java, "
    "ubicada en el package vetcare.util. La clase utiliza BufferedWriter (escritura) y BufferedReader "
    "(lectura) con codificacion UTF-8 para garantizar compatibilidad con caracteres especiales del espanol.")

add_body(doc, "Cada entidad se serializa en un archivo .txt independiente con campos separados por el "
    "delimitador '|'. Al iniciar el sistema, los datos se cargan automaticamente desde los archivos "
    "existentes en la carpeta data/. Al cerrar sesion, todos los datos se guardan automaticamente.")

add_bullet(doc, "duenos.txt        - id|nombre|dni|telefono|correo|direccion")
add_bullet(doc, "veterinarios.txt  - id|nombre|dni|telefono|correo|especialidad|colegiatura|activo")
add_bullet(doc, "mascotas.txt      - id|especie|nombre|raza|edad|sexo|peso|dnoDueno|observaciones")
add_bullet(doc, "consultas.txt     - idMascota|idConsulta|fecha|diagnostico|tratamiento|observaciones")

doc.add_paragraph()
add_subsection_title(doc, "10.3  Funcionalidades Implementadas")
add_body(doc, "Se implementaron 30 de los 40 requerimientos funcionales definidos, cubriendo los modulos "
    "de autenticacion, mascotas, duenos, veterinarios, citas e historial medico. Los 10 requerimientos "
    "restantes corresponden a funciones planificadas para versiones futuras del sistema (modificaciones "
    "individuales y reportes por periodo).", indent=False)

add_body(doc, "Estructura del proyecto en Eclipse (package vetcare.*):", indent=False)
add_bullet(doc, "vetcare.Main                       - Punto de entrada del sistema")
add_bullet(doc, "vetcare.gestion.SistemaVetCare     - Menu principal, autenticacion y flujo general")
add_bullet(doc, "vetcare.gestion.GestorMascotas     - CRUD de mascotas (implementa IGestionable<Mascota>)")
add_bullet(doc, "vetcare.gestion.GestorCitas        - CRUD de citas (implementa IGestionable<Cita>)")
add_bullet(doc, "vetcare.gestion.IGestionable<T>    - Interfaz generica de gestion")
add_bullet(doc, "vetcare.modelo.Animal / Perro / Gato / Ave - Jerarquia de herencia")
add_bullet(doc, "vetcare.modelo.Persona / Dueno / Veterinario - Jerarquia de personas")
add_bullet(doc, "vetcare.modelo.Mascota / Cita / ConsultaMedica - Entidades principales")
add_bullet(doc, "vetcare.util.PersistenciaArchivos  - Lectura y escritura de archivos .txt")

doc.add_page_break()

# ── SECCION 11: RESULTADOS ─────────────────────────────────────────────────────
add_section_title(doc, 11, "Resultados")
doc.add_paragraph()
add_body(doc, "Esta seccion presenta el estado de implementacion de los 40 requerimientos funcionales "
    "y las evidencias del funcionamiento del sistema VetCare.", indent=False)

add_subsection_title(doc, "11.1  Estado de Requerimientos Implementados")

tbl_estado = doc.add_table(rows=1, cols=4)
tbl_estado.style = 'Table Grid'
for i, h in enumerate(['RF', 'Nombre', 'Modulo', 'Estado']):
    tbl_estado.rows[0].cells[i].text = h
    run_h = tbl_estado.rows[0].cells[i].paragraphs[0].runs[0]
    run_h.font.bold = True; run_h.font.size = Pt(9.5); run_h.font.color.rgb = BLANCO
    set_cell_color(tbl_estado.rows[0].cells[i], '145A32')
    tbl_estado.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rf_estados = [
    ("RF-01", "Iniciar sesion",         "Autenticacion",    "Implementado"),
    ("RF-02", "Cerrar sesion",          "Autenticacion",    "Implementado"),
    ("RF-03", "Cambiar contrasena",     "Autenticacion",    "Planificado"),
    ("RF-04", "Registrar usuario",      "Autenticacion",    "Planificado"),
    ("RF-05", "Registrar mascota",      "Mascotas",         "Implementado"),
    ("RF-06", "Modificar mascota",      "Mascotas",         "Planificado"),
    ("RF-07", "Eliminar mascota",       "Mascotas",         "Implementado"),
    ("RF-08", "Buscar por nombre",      "Mascotas",         "Implementado"),
    ("RF-09", "Buscar por dueno",       "Mascotas",         "Implementado"),
    ("RF-10", "Listar mascotas",        "Mascotas",         "Implementado"),
    ("RF-11", "Ver ficha completa",     "Mascotas",         "Implementado"),
    ("RF-12", "Registrar especie",      "Mascotas",         "Implementado"),
    ("RF-13", "Registrar raza",         "Mascotas",         "Implementado"),
    ("RF-14", "Filtrar por especie",    "Mascotas",         "Implementado"),
    ("RF-15", "Registrar observaciones","Mascotas",         "Implementado"),
    ("RF-16", "Registrar dueno",        "Duenos",           "Implementado"),
    ("RF-17", "Modificar dueno",        "Duenos",           "Planificado"),
    ("RF-18", "Eliminar dueno",         "Duenos",           "Planificado"),
    ("RF-19", "Buscar dueno por nombre","Duenos",           "Planificado"),
    ("RF-20", "Buscar dueno por DNI",   "Duenos",           "Implementado"),
    ("RF-21", "Listar duenos",          "Duenos",           "Implementado"),
    ("RF-22", "Ver mascotas del dueno", "Duenos",           "Implementado"),
    ("RF-23", "Registrar cita",         "Citas",            "Implementado"),
    ("RF-24", "Modificar cita",         "Citas",            "Planificado"),
    ("RF-25", "Cancelar cita",          "Citas",            "Implementado"),
    ("RF-26", "Citas del dia",          "Citas",            "Implementado"),
    ("RF-27", "Buscar citas por mascota","Citas",           "Planificado"),
    ("RF-28", "Buscar citas por fecha", "Citas",            "Implementado"),
    ("RF-29", "Listar citas pendientes","Citas",            "Implementado"),
    ("RF-30", "Confirmar asistencia",   "Citas",            "Implementado"),
    ("RF-31", "Motivo de la cita",      "Citas",            "Implementado"),
    ("RF-32", "Registrar veterinario",  "Veterinarios",     "Implementado"),
    ("RF-33", "Modificar veterinario",  "Veterinarios",     "Planificado"),
    ("RF-34", "Desactivar veterinario", "Veterinarios",     "Implementado"),
    ("RF-35", "Listar veterinarios",    "Veterinarios",     "Implementado"),
    ("RF-36", "Ver agenda veterinario", "Veterinarios",     "Implementado"),
    ("RF-37", "Registrar consulta",     "Historial",        "Implementado"),
    ("RF-38", "Registrar tratamiento",  "Historial",        "Implementado"),
    ("RF-39", "Ver historial",          "Historial",        "Implementado"),
    ("RF-40", "Reporte por periodo",    "Historial",        "Planificado"),
]

for idx, (rf_id, nombre, modulo, estado) in enumerate(rf_estados):
    row = tbl_estado.add_row()
    bg = 'EAFAF1' if idx % 2 == 0 else 'FFFFFF'
    for i, txt in enumerate([rf_id, nombre, modulo, estado]):
        row.cells[i].text = txt
        run_d = row.cells[i].paragraphs[0].runs[0]
        run_d.font.size = Pt(9); run_d.font.color.rgb = GRIS_TEXTO
        if i == 0:
            run_d.font.bold = True; run_d.font.color.rgb = VERDE_OSCURO
        if i == 3:
            if estado == "Implementado":
                run_d.font.color.rgb = RGBColor(0x1A, 0x5E, 0x20)
                run_d.font.bold = True
            else:
                run_d.font.color.rgb = RGBColor(0x7D, 0x60, 0x08)
        set_cell_color(row.cells[i], bg)
    if estado == "Planificado":
        set_cell_color(row.cells[3], 'FEF9E7')

for i, w in enumerate([Cm(1.5), Cm(5.0), Cm(3.5), Cm(3.0)]):
    for cell in tbl_estado.columns[i].cells:
        cell.width = w

doc.add_paragraph()
add_body(doc, "Resumen: 30 requerimientos IMPLEMENTADOS | 10 requerimientos PLANIFICADOS (para versiones futuras)", indent=False)

doc.add_paragraph()
add_subsection_title(doc, "11.2  Evidencias del Sistema")
add_body(doc, "A continuacion se describen las pantallas principales del sistema VetCare. Las capturas "
    "de pantalla de la ejecucion en consola deben agregarse en esta seccion antes de la entrega final.")

evidencias = [
    "Pantalla de autenticacion: solicita usuario y contrasena con control de 3 intentos.",
    "Menu principal: 6 opciones de gestion (Mascotas, Duenos, Veterinarios, Citas, Historial, Citas del dia).",
    "Registro de mascota: ingreso de especie, datos del animal y seleccion de dueno por DNI.",
    "Listado de mascotas: muestra id, nombre, especie, raza, edad y dueno de cada mascota.",
    "Agendado de cita: seleccion de mascota, veterinario, fecha, hora y motivo; con validacion de conflictos.",
    "Historial medico: muestra consultas previas con fecha, diagnostico, tratamiento y observaciones.",
    "Citas del dia: lista automaticamente las citas agendadas para la fecha actual del sistema.",
]
for ev in evidencias:
    add_bullet(doc, ev)

add_body(doc, "NOTA: Agregar capturas de pantalla de la ejecucion en Eclipse antes de la entrega final.", indent=False)

doc.add_page_break()

# ── SECCION 12: CONCLUSIONES ───────────────────────────────────────────────────
add_section_title(doc, 12, "Conclusiones")
doc.add_paragraph()

conclusiones = [
    ("1.", "El sistema VetCare demuestra que la Programacion Orientada a Objetos permite modelar de "
       "manera natural y eficiente las entidades del dominio veterinario. La jerarquia de clases "
       "(Animal hacia Perro, Gato y Ave; Persona hacia Dueno y Veterinario) redujo la duplicacion de "
       "codigo y facilito la extension del sistema, validando el uso de herencia y polimorfismo como "
       "pilares del diseno."),
    ("2.", "La implementacion de 30 requerimientos funcionales de los 40 definidos evidencia que el "
       "sistema cubre el nucleo operativo de una clinica veterinaria: registro de pacientes, gestion "
       "de citas con control de conflictos, historial medico y persistencia de datos. Los 10 "
       "requerimientos restantes (modificaciones y reportes por periodo) constituyen la hoja de ruta "
       "para la siguiente iteracion del proyecto."),
    ("3.", "El uso de interfaces (IGestionable<T>) e internamente de ArrayList de Java permitio crear "
       "gestores de mascotas y citas reutilizables y extensibles. Esto confirma que el uso de "
       "colecciones genericas y contratos de interfaz es una practica solida en sistemas de gestion "
       "desarrollados con Java SE."),
    ("4.", "La persistencia mediante archivos .txt con PersistenciaArchivos.java resolvio el requisito "
       "de conservar datos entre sesiones sin depender de una base de datos externa, demostrando que "
       "es posible construir sistemas funcionales con las librerias estandar de Java (java.io) sin "
       "dependencias adicionales."),
    ("5.", "El uso de Git y GitHub como herramienta de control de versiones permitio mantener un "
       "historial completo del avance del proyecto, facilitar el trabajo incremental y generar "
       "evidencia objetiva del desarrollo realizado durante el ciclo academico."),
]

for num, texto in conclusiones:
    p_c = doc.add_paragraph()
    p_c.paragraph_format.left_indent = Cm(0.5)
    p_c.paragraph_format.space_after = Pt(6)
    run_num = p_c.add_run(f"{num}  ")
    run_num.font.bold = True; run_num.font.color.rgb = VERDE_OSCURO; run_num.font.size = Pt(11)
    run_txt = p_c.add_run(texto)
    run_txt.font.size = Pt(11); run_txt.font.color.rgb = GRIS_TEXTO

doc.add_page_break()

# ── SECCION 13: BIBLIOGRAFIA ───────────────────────────────────────────────────
add_section_title(doc, 13, "Bibliografia")
doc.add_paragraph()
add_body(doc, "Referencias bibliograficas en formato APA utilizadas como sustento del proyecto.", indent=False)
doc.add_paragraph()

referencias = [
    "Barrientos, J., & Molina, P. (2020). Sistema de informacion para la gestion de clinicas veterinarias. "
    "Repositorio Universidad de Guayaquil. http://repositorio.ug.edu.ec",

    "Deitel, P. J., & Deitel, H. M. (2012). Java: Como programar (9.a ed.). Pearson Educacion.",

    "Quispe, A., & Vargas, L. (2023). Desarrollo de software de escritorio para la gestion de consultas "
    "en clinicas veterinarias de Lima Metropolitana. Repositorio Institucional USIL. https://repositorio.usil.edu.pe",

    "Oracle Corporation. (2023). Java SE 17 Documentation. Oracle Help Center. "
    "https://docs.oracle.com/en/java/javase/17/",

    "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable "
    "Object-Oriented Software. Addison-Wesley Professional.",
]
for ref in referencias:
    p_r = doc.add_paragraph()
    p_r.paragraph_format.left_indent = Cm(0.5)
    p_r.paragraph_format.first_line_indent = Cm(-0.5)
    p_r.paragraph_format.space_after = Pt(6)
    run_r = p_r.add_run(ref)
    run_r.font.size = Pt(10.5); run_r.font.color.rgb = GRIS_TEXTO

doc.add_page_break()

# ── SECCION 14: ANEXOS ─────────────────────────────────────────────────────────
add_section_title(doc, 14, "Anexos")
doc.add_paragraph()

add_subsection_title(doc, "Anexo A - Enlace al Repositorio GitHub")
p_repo = doc.add_paragraph()
p_repo.paragraph_format.left_indent = Cm(0.5)
pPr_repo = p_repo._p.get_or_add_pPr()
shd_repo = OxmlElement('w:shd'); shd_repo.set(qn('w:fill'), 'EAFAF1'); shd_repo.set(qn('w:val'), 'clear')
pPr_repo.append(shd_repo)
r_repo = p_repo.add_run("Repositorio: https://github.com/webskechiza/Tecnicas-Programacion-OOP")
r_repo.font.size = Pt(10.5); r_repo.font.color.rgb = VERDE_MEDIO; r_repo.font.bold = True
add_body(doc, "Ruta del proyecto en el repositorio: Proyecto 2026/Vetcare/Desarrollo/")
add_body(doc, "El repositorio contiene el historial completo de commits realizados durante el ciclo 2026-I, "
    "incluyendo las versiones parciales (VetCare_Avance) y la version completa (Desarrollo/).")

doc.add_paragraph()
add_subsection_title(doc, "Anexo B - Estructura de Packages del Codigo Fuente")

codigo_estructura = """package vetcare;           -> Main.java
package vetcare.gestion;   -> IGestionable.java
                              GestorMascotas.java
                              GestorCitas.java
                              SistemaVetCare.java
package vetcare.modelo;    -> Animal.java   (abstracta)
                              Perro.java    (extends Animal)
                              Gato.java     (extends Animal)
                              Ave.java      (extends Animal)
                              Persona.java  (abstracta)
                              Dueno.java    (extends Persona)
                              Veterinario.java (extends Persona)
                              Mascota.java
                              Cita.java
                              ConsultaMedica.java
package vetcare.util;      -> PersistenciaArchivos.java"""

p_cod = doc.add_paragraph()
p_cod.paragraph_format.left_indent = Cm(0.5)
pPr_cod = p_cod._p.get_or_add_pPr()
shd_cod = OxmlElement('w:shd'); shd_cod.set(qn('w:fill'), '2C3E50'); shd_cod.set(qn('w:val'), 'clear')
pPr_cod.append(shd_cod)
r_cod = p_cod.add_run(codigo_estructura)
r_cod.font.name = 'Courier New'; r_cod.font.size = Pt(9); r_cod.font.color.rgb = AMBAR_CLARO

doc.add_paragraph()
add_subsection_title(doc, "Anexo C - Fragmento de Codigo: Main.java")

main_codigo = """package vetcare;

import vetcare.gestion.SistemaVetCare;

public class Main {
    public static void main(String[] args) {
        SistemaVetCare sistema = new SistemaVetCare("data");
        sistema.iniciar();
    }
}"""

p_main = doc.add_paragraph()
p_main.paragraph_format.left_indent = Cm(0.5)
pPr_main = p_main._p.get_or_add_pPr()
shd_main = OxmlElement('w:shd'); shd_main.set(qn('w:fill'), '2C3E50'); shd_main.set(qn('w:val'), 'clear')
pPr_main.append(shd_main)
r_main = p_main.add_run(main_codigo)
r_main.font.name = 'Courier New'; r_main.font.size = Pt(9); r_main.font.color.rgb = AMBAR_CLARO

doc.add_paragraph()
add_subsection_title(doc, "Anexo D - Fragmento de Codigo: IGestionable.java")

interfaz_codigo = """package vetcare.gestion;

import java.util.List;

public interface IGestionable<T> {
    void agregar(T elemento);
    T buscarPorId(int id);
    List<T> getTodas();
    boolean eliminar(int id);
    void listar();
}"""

p_intf = doc.add_paragraph()
p_intf.paragraph_format.left_indent = Cm(0.5)
pPr_intf = p_intf._p.get_or_add_pPr()
shd_intf = OxmlElement('w:shd'); shd_intf.set(qn('w:fill'), '2C3E50'); shd_intf.set(qn('w:val'), 'clear')
pPr_intf.append(shd_intf)
r_intf = p_intf.add_run(interfaz_codigo)
r_intf.font.name = 'Courier New'; r_intf.font.size = Pt(9); r_intf.font.color.rgb = AMBAR_CLARO

# ── PIE DE PAGINA ──────────────────────────────────────────────────────────────
doc.add_page_break()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_f = p_footer._p.get_or_add_pPr()
shd_f = OxmlElement('w:shd'); shd_f.set(qn('w:fill'), '145A32'); shd_f.set(qn('w:val'), 'clear')
pPr_f.append(shd_f)
r_f = p_footer.add_run("  VetCare - Sistema de Gestion de Clinica Veterinaria  |  UPN 2026-I  "
    "|  Tecnicas de Programacion Orientada a Objetos  ")
r_f.font.size = Pt(9); r_f.font.color.rgb = AMBAR_CLARO

# ── GUARDAR ────────────────────────────────────────────────────────────────────
out_path = os.path.join(DOC_DIR, "VetCare_InformeFinal.docx")
doc.save(out_path)
print(f"Documento guardado en: {out_path}")
print("Abre el archivo en Word y ejecuta el corrector ortografico (F7) antes de entregar.")
